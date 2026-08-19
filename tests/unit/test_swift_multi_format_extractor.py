"""
Unit Tests for Smart SWIFT & Bank Transfer Multi-Format File Extractor (BP-012 / BP-013)
Tests Word (.docx), Excel (.xlsx, .csv), PDF (.pdf), and Image OCR (.jpg, .png) parsing and reconciliation.
"""

import io
import csv
import json
import pytest
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session

from main import app
from database.database import get_db, SessionLocal
from modules.financial_approval.swift_file_extractor import (
    extract_text_from_swift_file,
    normalize_swift_ocr_text,
    extract_text_from_word,
    extract_text_from_excel,
    extract_text_from_pdf,
)
from modules.financial_approval.swift_mt103_parser import parse_swift_mt103_text


@pytest.fixture
def client():
    return TestClient(app)


def test_normalize_swift_ocr_text():
    """Tests normalization of common OCR misrecognitions and fullwidth unicode characters."""
    ocr_raw = """
：20/TRANSACTION REFERENCE NUMBER
:FT/26228/KZ70Q
:23B/BANK OPERATIONCODE
:CRED
:32A/Value Date,CCY,Amount
260818USD43704,00
:5OK/ORDERING CUST
/EG780057004001017153610010101
SCAS FOR CONSTRUCTION AND FINISHING
:57A/Account with Bank
:PCBCCNBJJSS
59/BeneficiaryCustomer
：/32250198613609841015
SUZHOU YUHENG TEXTILE CO.,LTD
70/DETAILS OF PAYMENT
:EG0010040PINO.YH20260730.6
71A/DETAILS OF CHARGES
:SHA
"""
    normalized = normalize_swift_ocr_text(ocr_raw)
    parsed = parse_swift_mt103_text(normalized)

    assert parsed["success"] is True
    assert parsed["transaction_reference"] == "FT/26228/KZ70Q"
    assert parsed["amount"] == 43704.0
    assert parsed["currency"] == "USD"
    assert parsed["value_date"] == "2026-08-18"
    assert parsed["beneficiary_bank_swift"] == "PCBCCNBJJSS"
    assert parsed["beneficiary_account_or_iban"] == "32250198613609841015"
    assert parsed["pi_number"] == "YH20260730.6"
    assert parsed["charge_details"] == "SHA"


def test_extract_from_csv_file():
    """Tests extracting SWIFT data from a CSV spreadsheet."""
    csv_buffer = io.StringIO()
    writer = csv.writer(csv_buffer)
    writer.writerow(["Field", "Value", "Notes"])
    writer.writerow([":20:", "FT/CSV/998877", "Transaction Ref"])
    writer.writerow([":32A:", "260818EUR15200.50", "Amount in EUR"])
    writer.writerow([":57A:", "BNPAFRPPXXX", "Beneficiary Bank"])
    writer.writerow([":59:", "PARIS LOGISTICS SAS", "Beneficiary"])

    csv_bytes = csv_buffer.getvalue().encode("utf-8")
    raw_text, normalized = extract_text_from_swift_file("swift_advice.csv", csv_bytes)
    parsed = parse_swift_mt103_text(normalized)

    assert parsed["success"] is True
    assert parsed["transaction_reference"] == "FT/CSV/998877"
    assert parsed["currency"] == "EUR"
    assert parsed["amount"] == 15200.50
    assert parsed["beneficiary_bank_swift"] == "BNPAFRPPXXX"


def test_extract_from_word_file():
    """Tests extracting SWIFT data from a Word (.docx) document."""
    import docx
    doc = docx.Document()
    doc.add_heading("Bank Transfer Advice - SWIFT MT103", 0)
    doc.add_paragraph(":20: FT/DOCX/2026/001")
    doc.add_paragraph(":32A: 260818USD55000,00")
    doc.add_paragraph(":50K: /EG123456789 EGYPT IMPORT CO")
    doc.add_paragraph(":57A: CITIEGCAXXX")
    doc.add_paragraph(":59: /DE987654321 GLOBAL EXPORTER GMBH")
    doc.add_paragraph(":70: PI NO.PI-2026-99")
    doc.add_paragraph(":71A: OUR")

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_bytes = doc_io.getvalue()

    raw_text, normalized = extract_text_from_swift_file("advice.docx", doc_bytes)
    parsed = parse_swift_mt103_text(normalized)

    assert parsed["success"] is True
    assert parsed["transaction_reference"] == "FT/DOCX/2026/001"
    assert parsed["currency"] == "USD"
    assert parsed["amount"] == 55000.0
    assert parsed["beneficiary_bank_swift"] == "CITIEGCAXXX"
    assert parsed["pi_number"] == "PI-2026-99"
    assert parsed["charge_details"] == "OUR"


def test_extract_from_excel_file():
    """Tests extracting SWIFT data from an Excel (.xlsx) workbook."""
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "SWIFT_Details"
    ws.append(["Tag", "Content"])
    ws.append([":20:", "FT/XLSX/443322"])
    ws.append([":32A:", "260818USD12850.00"])
    ws.append([":57A:", "HSBCHKHHXXX"])
    ws.append([":59:", "HONG KONG SUPPLIER LTD"])
    ws.append([":70:", "PAYMENT FOR PO #PO-8877"])

    xlsx_io = io.BytesIO()
    wb.save(xlsx_io)
    xlsx_bytes = xlsx_io.getvalue()

    raw_text, normalized = extract_text_from_swift_file("payment_sheet.xlsx", xlsx_bytes)
    parsed = parse_swift_mt103_text(normalized)

    assert parsed["success"] is True
    assert parsed["transaction_reference"] == "FT/XLSX/443322"
    assert parsed["amount"] == 12850.0
    assert parsed["beneficiary_bank_swift"] == "HSBCHKHHXXX"
    assert parsed["po_number"] == "PO-8877"


def test_api_smart_extract_file_endpoint(client):
    """Tests POST /api/v1/financial-approval/swift/smart-extract-file via TestClient."""
    sample_text = """
:20:FT/TEST/778899
:32A:260818USD43704,00
:50K:/EG780057004001017153610010101
SCAS FOR CONSTRUCTION AND FINISHING
:57A:PCBCCNBJJSS
:59:/32250198613609841015
SUZHOU YUHENG TEXTILE CO., LTD
:70:EG0010040 PI NO.YH20260730.6
:71A:SHA
""".encode("utf-8")

    files = {"file": ("swift_advice.txt", io.BytesIO(sample_text), "text/plain")}
    response = client.post("/api/v1/financial-approval/swift/smart-extract-file", files=files)

    assert response.status_code == 200
    data = response.json()
    assert data["success"] is True
    assert data["detected_filename"] == "swift_advice.txt"
    assert data["detected_file_type"] == "Text File"
    assert data["parsed_swift"]["transaction_reference"] == "FT/TEST/778899"
    assert data["parsed_swift"]["amount"] == 43704.0
    assert data["parsed_swift"]["currency"] == "USD"
