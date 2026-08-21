"""
Service Layer & Business Engine for Import Documentation & ACI (Phase 3 - BP-014 to BP-019)
"""

import io
import re
from datetime import date, datetime, timezone
from sqlalchemy.orm import Session
from fastapi import HTTPException, status

from modules.import_documentation.model import (
    AcidRegistrationSession,
    BankingDocumentSession,
    ShipmentDocumentItem,
    CustomsDeclarationDraft,
    POPackingReconciliationSession,
    DraftBLReviewSession,
    CertificateOfOriginReviewSession,
    InspectionCertificateReviewSession,
)
from modules.import_documentation.schemas import (
    AcidRegistrationCreate,
    AcidRegistrationUpdate,
    AcidRegistrationResponse,
    BankingDocumentCreate,
    ShipmentDocumentCreate,
    ShipmentDocumentUpdate,
    CustomsDeclarationCreate,
    POReconciliationSessionCreate,
    POReconciliationSessionUpdate,
    POReconciliationSessionResponse,
    COODraftTemplateResponse,
    COOComparisonRequest,
    InspectionDraftTemplateResponse,
    InspectionComparisonRequest,
    DocumentExtractRequest,
    DocumentExtractResponse,
    ThreeWayCrossMatchRequest,
    ThreeWayCrossMatchResponse,
    CentralArchiveDocumentSummary,
    CentralArchiveResponse,
)

from modules.import_documentation.validators import (
    validate_acid_number,
    validate_acid_expiry,
    validate_no_duplicate_acid_session,
    validate_no_duplicate_form4_session,
)
import modules.import_documentation.repository as repo
from modules.customs_tariff.model import CustomsTariff


from modules.import_documentation.nafeza_acid_parser import (
    parse_nafeza_acid_text,
    compare_acid_data,
    generate_whatsapp_request_text,
    generate_email_request_template,
)
from modules.import_documentation.ai_document_parser import (
    extract_coo_china_ccpit_text,
    extract_eur1_certificate_text,
    extract_inspection_voc_certificate_text,
    _heuristic_multi_carrier_extractor as extract_draft_bl_data,
    extract_commercial_invoice_data,
)


def enrich_acid_response(db: Session, item: AcidRegistrationSession, import_files_map: dict = None) -> AcidRegistrationResponse:
    today = date.today()
    # Safely compute days remaining — expiry_date can be None for legacy/pending records
    if item.expiry_date:
        try:
            days_rem = (item.expiry_date - today).days
        except Exception:
            days_rem = 0
    else:
        days_rem = 0
    is_ver = (
        item.status in ["Verified", "Discrepancy_Accepted"]
        and item.is_importer_matched
        and item.is_exporter_matched
        and item.is_invoice_matched
        and item.is_ports_matched
        and days_rem > 0
    )

    import_file_code = None
    if item.import_file_id:
        if import_files_map is not None:
            imp = import_files_map.get(item.import_file_id)
            if imp:
                import_file_code = imp.import_file_code or imp.custom_file_number
        else:
            from modules.import_files.model import ImportFile
            imp = db.query(ImportFile).filter(ImportFile.import_file_id == item.import_file_id).first()
            if imp:
                import_file_code = imp.import_file_code or imp.custom_file_number

    return AcidRegistrationResponse(
        acid_id=item.acid_id,
        acid_code=item.acid_code,
        acid_number=item.acid_number,
        import_file_id=item.import_file_id,
        import_file_code=import_file_code,
        po_id=item.po_id,
        po_number=item.po_number,
        po_date=item.po_date,
        importer_id=item.importer_id,
        importer_name=item.importer_name,
        importer_tax_id=item.importer_tax_id,
        importer_address=item.importer_address,
        supplier_id=item.supplier_id,
        exporter_name=item.exporter_name,
        exporter_reg_type=item.exporter_reg_type,
        exporter_reg_id=item.exporter_reg_id,
        exporter_country=item.exporter_country,
        exporter_country_code=item.exporter_country_code,
        exporter_address=item.exporter_address,
        exporter_phone=item.exporter_phone,
        cargox_id=item.cargox_id,
        proforma_invoice_no=item.proforma_invoice_no,
        proforma_invoice_date=item.proforma_invoice_date,
        invoice_date=item.invoice_date,
        invoice_type=item.invoice_type,
        invoice_attachment_name=item.invoice_attachment_name,
        pol_name=item.pol_name,
        pod_name=item.pod_name,
        customs_broker_id=item.customs_broker_id,
        customs_broker_name=item.customs_broker_name,
        customs_broker_phone=item.customs_broker_phone,
        requested_date=item.requested_date,
        generated_date=item.generated_date,
        expiry_date=item.expiry_date,
        raw_nafeza_text=item.raw_nafeza_text,
        requested_data=item.requested_data,
        generated_data=item.generated_data,
        discrepancies_data=item.discrepancies_data,
        discrepancy_override_reason=item.discrepancy_override_reason,
        is_importer_matched=item.is_importer_matched,
        is_exporter_matched=item.is_exporter_matched,
        is_invoice_matched=item.is_invoice_matched,
        is_ports_matched=item.is_ports_matched,
        has_discrepancies=item.has_discrepancies,
        verification_notes=item.verification_notes,
        status=item.status,
        days_to_expiry=days_rem,
        execution_days=item.execution_days,
        is_verified=is_ver,
        is_active=item.is_active,
        created_at=item.created_at,
        updated_at=item.updated_at,
    )


def enrich_banking_response(db: Session, item: BankingDocumentSession):
    import_file_code = None
    importer_name = None
    supplier_name = None
    po_number = None

    if item.import_file_id:
        from modules.import_files.model import ImportFile
        imp = db.query(ImportFile).filter(ImportFile.import_file_id == item.import_file_id).first()
        if imp:
            import_file_code = imp.import_file_code or imp.custom_file_number
            importer_name = imp.company_name
            supplier_name = imp.supplier_name
            po_number = imp.po_number

    from modules.import_documentation.schemas import BankingDocumentResponse
    res = BankingDocumentResponse.model_validate(item)
    res.import_file_code = import_file_code
    res.importer_name = importer_name
    res.supplier_name = supplier_name
    res.po_number = po_number
    return res


def enrich_shipment_doc_response(db: Session, item: ShipmentDocumentItem):
    import_file_code = None
    if item.import_file_id:
        from modules.import_files.model import ImportFile
        imp = db.query(ImportFile).filter(ImportFile.import_file_id == item.import_file_id).first()
        if imp:
            import_file_code = imp.import_file_code or imp.custom_file_number

    from modules.import_documentation.schemas import ShipmentDocumentResponse
    res = ShipmentDocumentResponse.model_validate(item)
    res.import_file_code = import_file_code
    return res


def create_acid_session_service(
    db: Session, schema: AcidRegistrationCreate
) -> AcidRegistrationResponse:
    validate_acid_number(schema.acid_number, allow_pending=True)
    validate_acid_expiry(schema.expiry_date, schema.requested_date)

    # Check if an active session already exists for this import file - if so, seamlessly update it (Upsert)
    if schema.import_file_id:
        from modules.import_documentation.model import AcidRegistrationSession
        existing = (
            db.query(AcidRegistrationSession)
            .filter(
                AcidRegistrationSession.import_file_id == schema.import_file_id,
                AcidRegistrationSession.is_active == True,
            )
            .first()
        )
        if existing:
            update_data = schema.model_dump(exclude_unset=True)
            for k, v in update_data.items():
                if hasattr(existing, k) and v is not None:
                    setattr(existing, k, v)
            if hasattr(existing, "set_updated_info"):
                existing.set_updated_info()
            
            # Compute execution days if dates are present
            if existing.generated_date and existing.requested_date:
                existing.execution_days = max(0, (existing.generated_date - existing.requested_date).days)

            db.commit()
            db.refresh(existing)
            db_item = existing

            # Sync with import file if applicable
            if db_item.import_file_id and db_item.acid_number and db_item.acid_number != "PENDING":
                from modules.import_files.model import ImportFile
                imp = db.query(ImportFile).filter(ImportFile.import_file_id == db_item.import_file_id).first()
                if imp:
                    imp.acid_number = db_item.acid_number
                    imp.acid_request_date = db_item.requested_date
                    imp.acid_issue_date = db_item.generated_date or db_item.requested_date
                    imp.acid_expiry_date = db_item.expiry_date
                    imp.acid_execution_days = db_item.execution_days
                    db.commit()

            return enrich_acid_response(db, db_item)

    db_item = repo.create_acid_session(db, schema)

    # Compute execution days if dates are present
    if db_item.generated_date and db_item.requested_date:
        db_item.execution_days = max(0, (db_item.generated_date - db_item.requested_date).days)
        db.commit()

    # Sync with import file if applicable
    if db_item.import_file_id and db_item.acid_number and db_item.acid_number != "PENDING":
        from modules.import_files.model import ImportFile
        imp = db.query(ImportFile).filter(ImportFile.import_file_id == db_item.import_file_id).first()
        if imp:
            imp.acid_number = db_item.acid_number
            imp.acid_request_date = db_item.requested_date
            imp.acid_issue_date = db_item.generated_date or db_item.requested_date
            imp.acid_expiry_date = db_item.expiry_date
            imp.acid_execution_days = db_item.execution_days
            db.commit()

    return enrich_acid_response(db, db_item)


def update_acid_session_service(
    db: Session, acid_id: int, schema: AcidRegistrationUpdate
) -> AcidRegistrationResponse:
    db_item = repo.get_acid_session_by_id(db, acid_id)
    if not db_item:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"ACID Registration Session ID {acid_id} not found.",
        )

    if schema.import_file_id is not None:
        validate_no_duplicate_acid_session(db, schema.import_file_id, current_acid_id=acid_id)

    if schema.acid_number:
        validate_acid_number(schema.acid_number, allow_pending=True)

    updated = repo.update_acid_session(db, db_item, schema)

    # Compute execution days if dates are present
    if updated.generated_date and updated.requested_date:
        updated.execution_days = max(0, (updated.generated_date - updated.requested_date).days)
        db.commit()

    # Sync with import file if applicable
    if updated.import_file_id and updated.acid_number and updated.acid_number != "PENDING":
        from modules.import_files.model import ImportFile
        imp = db.query(ImportFile).filter(ImportFile.import_file_id == updated.import_file_id).first()
        if imp:
            imp.acid_number = updated.acid_number
            if updated.requested_date:
                imp.acid_request_date = updated.requested_date
            if updated.generated_date:
                imp.acid_issue_date = updated.generated_date
            elif updated.requested_date and not imp.acid_issue_date:
                imp.acid_issue_date = updated.requested_date
            if updated.expiry_date:
                imp.acid_expiry_date = updated.expiry_date
            if updated.execution_days is not None:
                imp.acid_execution_days = updated.execution_days
            db.commit()

    return enrich_acid_response(db, updated)


def get_acid_tracker_service(db: Session):
    from datetime import timedelta
    from modules.import_files.model import ImportFile
    from modules.customs_clearance.model import CustomsClearanceRecord
    from modules.import_documentation.schemas import AcidTrackerItem, AcidTrackerSummary

    today = date.today()
    items = []

    # Map clearances for rapid release check
    clearances = db.query(CustomsClearanceRecord).filter(CustomsClearanceRecord.is_active == True).all()
    clearance_by_file = {}
    for cl in clearances:
        if cl.import_file_id not in clearance_by_file:
            clearance_by_file[cl.import_file_id] = cl
        elif cl.dispatch_authorized or cl.release_permit_no:
            clearance_by_file[cl.import_file_id] = cl

    # Query all acid sessions
    acid_sessions = db.query(AcidRegistrationSession).filter(AcidRegistrationSession.is_active == True).all()
    acid_sessions_by_file = {s.import_file_id: s for s in acid_sessions if s.import_file_id}
    processed_session_ids = set()

    # Query all import files
    import_files = db.query(ImportFile).filter(ImportFile.is_active == True).all()

    for imp in import_files:
        acid_sess = acid_sessions_by_file.get(imp.import_file_id)
        if acid_sess:
            processed_session_ids.add(acid_sess.acid_id)

        acid_num = imp.acid_number or (acid_sess.acid_number if acid_sess else None)
        if not acid_num and not acid_sess:
            continue

        acid_issue_date = imp.acid_issue_date or (acid_sess.generated_date if acid_sess else None) or (acid_sess.requested_date if acid_sess else None)
        acid_expiry_date = imp.acid_expiry_date or (acid_sess.expiry_date if acid_sess else None)
        if not acid_expiry_date and acid_issue_date:
            acid_expiry_date = acid_issue_date + timedelta(days=90)

        cl = clearance_by_file.get(imp.import_file_id)
        is_released = (
            imp.is_customs_released
            or (cl is not None and (cl.dispatch_authorized or bool(cl.release_permit_no) or cl.status in ["Final Release Granted", "Cleared", "Completed"]))
            or imp.current_stage in ["Phase 8 - Warehouse Receiving", "Phase 9 - Financial Settlement", "Phase 10 - Closure", "Warehouse Received", "Closed"]
        )
        customs_released_at = imp.customs_released_at or (cl.release_date if cl else None) or (cl.dispatch_date if cl else None)
        release_permit_no = cl.release_permit_no if cl else None

        days_rem = (acid_expiry_date - today).days if acid_expiry_date else 90
        total_days = (acid_expiry_date - acid_issue_date).days if (acid_expiry_date and acid_issue_date) else 90
        if total_days <= 0:
            total_days = 90
        validity_pct = max(0.0, min(100.0, (days_rem / total_days) * 100.0)) if not is_released else 100.0

        if is_released:
            status = "Customs Released"
            status_label_ar = "صُرفت من الجمرك (معفى من التنبيه)"
            alert_required = False
        elif not acid_num or acid_num.upper() in ["PENDING", "REQUESTED", "DRAFT", ""]:
            status = "Pending Issue"
            status_label_ar = "قيد استخراج الـ ACID"
            alert_required = False
        elif days_rem <= 0:
            status = "Expired"
            status_label_ar = f"منتهي الصلاحية (منذ {abs(days_rem)} يوم)"
            alert_required = True
        elif days_rem <= 14:
            status = "Expiring Soon"
            status_label_ar = f"يوشك على الانتهاء (متبقي {days_rem} يوم)"
            alert_required = True
        else:
            status = "Valid"
            status_label_ar = f"ساري وصالح (متبقي {days_rem} يوم)"
            alert_required = False

        items.append(
            AcidTrackerItem(
                import_file_id=imp.import_file_id,
                import_file_code=imp.import_file_code or imp.custom_file_number,
                custom_file_number=imp.custom_file_number,
                acid_session_id=acid_sess.acid_id if acid_sess else None,
                acid_code=acid_sess.acid_code if acid_sess else None,
                acid_number=acid_num or "PENDING",
                importer_name=imp.company_name or (acid_sess.importer_name if acid_sess else "الشركة المستوردة"),
                supplier_name=imp.supplier_name or (acid_sess.exporter_name if acid_sess else "المورد الأجنبي"),
                po_number=imp.po_number or (acid_sess.po_number if acid_sess else None),
                pi_number=imp.pi_number or (acid_sess.proforma_invoice_no if acid_sess else None),
                shipment_mode=imp.shipment_mode,
                current_stage=imp.current_stage,
                customs_broker_name=imp.broker_name or (acid_sess.customs_broker_name if acid_sess else None),
                acid_issue_date=acid_issue_date,
                acid_expiry_date=acid_expiry_date,
                execution_days=imp.acid_execution_days or (acid_sess.execution_days if acid_sess else None),
                days_remaining=days_rem,
                total_validity_days=total_days,
                validity_percentage=round(validity_pct, 1),
                is_customs_released=is_released,
                customs_released_at=customs_released_at,
                release_permit_no=release_permit_no,
                status=status,
                status_label_ar=status_label_ar,
                alert_required=alert_required,
            )
        )

    # Standalone acid sessions
    for sess in acid_sessions:
        if sess.acid_id in processed_session_ids:
            continue
        days_rem = (sess.expiry_date - today).days if sess.expiry_date else 90
        total_days = (sess.expiry_date - (sess.generated_date or sess.requested_date)).days if sess.expiry_date else 90
        if total_days <= 0:
            total_days = 90
        validity_pct = max(0.0, min(100.0, (days_rem / total_days) * 100.0))

        if sess.acid_number == "PENDING":
            status = "Pending Issue"
            status_label_ar = "قيد استخراج الـ ACID"
            alert_required = False
        elif days_rem <= 0:
            status = "Expired"
            status_label_ar = f"منتهي الصلاحية (منذ {abs(days_rem)} يوم)"
            alert_required = True
        elif days_rem <= 14:
            status = "Expiring Soon"
            status_label_ar = f"يوشك على الانتهاء (متبقي {days_rem} يوم)"
            alert_required = True
        else:
            status = "Valid"
            status_label_ar = f"ساري وصالح (متبقي {days_rem} يوم)"
            alert_required = False

        items.append(
            AcidTrackerItem(
                import_file_id=None,
                import_file_code=None,
                custom_file_number=None,
                acid_session_id=sess.acid_id,
                acid_code=sess.acid_code,
                acid_number=sess.acid_number,
                importer_name=sess.importer_name,
                supplier_name=sess.exporter_name,
                po_number=sess.po_number,
                pi_number=sess.proforma_invoice_no,
                shipment_mode=None,
                current_stage="ACID Registration",
                customs_broker_name=sess.customs_broker_name,
                acid_issue_date=sess.generated_date or sess.requested_date,
                acid_expiry_date=sess.expiry_date,
                execution_days=sess.execution_days,
                days_remaining=days_rem,
                total_validity_days=total_days,
                validity_percentage=round(validity_pct, 1),
                is_customs_released=False,
                customs_released_at=None,
                release_permit_no=None,
                status=status,
                status_label_ar=status_label_ar,
                alert_required=alert_required,
            )
        )

    valid_count = sum(1 for i in items if i.status == "Valid")
    expiring_soon_count = sum(1 for i in items if i.status == "Expiring Soon")
    expired_count = sum(1 for i in items if i.status == "Expired")
    customs_released_count = sum(1 for i in items if i.status == "Customs Released")
    pending_issue_count = sum(1 for i in items if i.status == "Pending Issue")

    return AcidTrackerSummary(
        total_acids_count=len(items),
        valid_count=valid_count,
        expiring_soon_count=expiring_soon_count,
        expired_count=expired_count,
        customs_released_count=customs_released_count,
        pending_issue_count=pending_issue_count,
        items=items,
    )



def parse_acid_text_service(
    db: Session, raw_text: str, import_file_id: int | None = None
) -> dict:
    parsed = parse_nafeza_acid_text(raw_text)
    comparison = None

    if import_file_id:
        from modules.import_files.model import ImportFile
        from modules.suppliers.model import Supplier
        from modules.import_companies.model import ImportCompany

        file_obj = db.query(ImportFile).filter(ImportFile.import_file_id == import_file_id).first()
        if file_obj:
            comp_obj = db.query(ImportCompany).filter(ImportCompany.company_id == file_obj.company_id).first() if file_obj.company_id else None
            supp_obj = db.query(Supplier).filter(Supplier.supplier_id == file_obj.supplier_id).first() if file_obj.supplier_id else None

            requested_dict = {
                "importer_name": comp_obj.importer_name if comp_obj else file_obj.company_name,
                "importer_tax_id": comp_obj.vat_id if comp_obj else "",
                "exporter_name": supp_obj.company_name if supp_obj else file_obj.supplier_name,
                "exporter_reg_type": supp_obj.registration_type if supp_obj else "VAT Number",
                "exporter_reg_id": supp_obj.foreign_exporter_id if supp_obj else "",
                "exporter_country": supp_obj.foreign_exporter_country if supp_obj else "",
                "exporter_country_code": supp_obj.foreign_exporter_country_code if supp_obj else "",
                "proforma_invoice_no": file_obj.pi_number or "",
                "pol_name": "",
                "pod_name": "",
                "cargox_id": getattr(supp_obj, "cargox_platform_id", None) or "",
            }
            comparison = compare_acid_data(requested_dict, parsed)

    return {
        "parsed_data": parsed,
        "comparison": comparison,
    }


def compare_acid_datasets_service(requested: dict, generated: dict) -> dict:
    return compare_acid_data(requested, generated)


def generate_acid_templates_service(req_data: dict) -> dict:
    wa_text = generate_whatsapp_request_text(req_data)
    email_dict = generate_email_request_template(req_data)
    return {
        "whatsapp_text": wa_text,
        "email_subject": email_dict["subject"],
        "email_body": email_dict["body"],
    }


def restore_acid_session_service(db: Session, acid_id: int) -> AcidRegistrationResponse:
    item = repo.restore_acid_session(db, acid_id)
    if not item:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"ACID Registration Session ID {acid_id} not found.",
        )
    return enrich_acid_response(db, item)


# --- BANKING DOCUMENTS SERVICE ---
def create_banking_document_service(
    db: Session, schema: BankingDocumentCreate
):
    if schema.doc_type == "Form 4":
        validate_no_duplicate_form4_session(db, schema.import_file_id)

    item = repo.create_banking_document(db, schema)
    
    # Sync with import file if applicable
    if item.import_file_id:
        from modules.import_files.model import ImportFile
        imp = db.query(ImportFile).filter(ImportFile.import_file_id == item.import_file_id).first()
        if imp:
            imp.form4_request_date = item.request_date
            if item.doc_reference_number and item.doc_reference_number != "PENDING":
                imp.form4_no = item.doc_reference_number
            if item.received_date:
                imp.form4_received_date = item.received_date
                imp.form4_execution_days = item.execution_days
            db.commit()

    return enrich_banking_response(db, item)


def update_banking_document_service(
    db: Session, bank_doc_id: int, schema: BankingDocumentUpdate
):
    item = repo.get_banking_document_by_id(db, bank_doc_id)
    if not item:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Banking Document ID {bank_doc_id} not found.",
        )

    if schema.import_file_id is not None:
        validate_no_duplicate_form4_session(db, schema.import_file_id, current_doc_id=bank_doc_id)

    updated = repo.update_banking_document(db, item, schema)

    # Sync with import file if applicable
    if updated.import_file_id:
        from modules.import_files.model import ImportFile
        imp = db.query(ImportFile).filter(ImportFile.import_file_id == updated.import_file_id).first()
        if imp:
            if updated.doc_reference_number and updated.doc_reference_number != "PENDING":
                imp.form4_no = updated.doc_reference_number
            if updated.request_date:
                imp.form4_request_date = updated.request_date
            if updated.received_date:
                imp.form4_received_date = updated.received_date
                imp.form4_execution_days = updated.execution_days
            db.commit()

    return enrich_banking_response(db, updated)


def receive_banking_document_service(
    db: Session, bank_doc_id: int, form4_number: str, received_date: date, notes: str | None = None
):
    item = repo.get_banking_document_by_id(db, bank_doc_id)
    if not item:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Banking Document ID {bank_doc_id} not found.",
        )

    req_date = item.request_date or item.issue_date or date.today()
    exec_days = max(0, (received_date - req_date).days)

    item.doc_reference_number = form4_number.strip()
    item.received_date = received_date
    item.execution_days = exec_days
    item.status = "Received"
    if notes:
        item.notes = f"{item.notes}\n{notes}" if item.notes else notes

    item.updated_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(item)

    # Sync with import file
    if item.import_file_id:
        from modules.import_files.model import ImportFile
        imp = db.query(ImportFile).filter(ImportFile.import_file_id == item.import_file_id).first()
        if imp:
            imp.form4_no = item.doc_reference_number
            imp.form4_request_date = item.request_date
            imp.form4_received_date = item.received_date
            imp.form4_execution_days = item.execution_days
            db.commit()

    return enrich_banking_response(db, item)


def delete_banking_document_service(db: Session, bank_doc_id: int):
    item = repo.get_banking_document_by_id(db, bank_doc_id)
    if not item:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Banking Document ID {bank_doc_id} not found.",
        )
    repo.delete_banking_document(db, item)
    return {"status": "success", "message": f"Banking Document {bank_doc_id} deleted"}


# --- SHIPMENT DOCUMENTS SERVICE ---
def create_shipment_document_service(
    db: Session, schema: ShipmentDocumentCreate
):
    item = repo.create_shipment_document(db, schema)
    return enrich_shipment_doc_response(db, item)


def update_cargox_and_bl_endorsement_service(
    db: Session,
    doc_id: int,
    cargox_envelope_id: str | None = None,
    endorsement_number: str | None = None,
):
    schema = ShipmentDocumentUpdate()
    if cargox_envelope_id:
        schema.is_cargox_uploaded = True
        schema.cargox_envelope_id = cargox_envelope_id
    if endorsement_number:
        schema.is_bl_endorsed = True
        schema.endorsement_number = endorsement_number
        schema.status = "Endorsed"

    item = repo.update_shipment_document(db, doc_id, schema)
    return enrich_shipment_doc_response(db, item)


# --- CUSTOMS DECLARATION SERVICE ---
def create_customs_declaration_service(
    db: Session, schema: CustomsDeclarationCreate
) -> CustomsDeclarationDraft:
    return repo.create_customs_declaration(db, schema)


# ==============================================================================
# PHASE 6: INTELLIGENT DOCUMENT VERIFICATION & COMPARISON ENGINES
# ==============================================================================
import difflib
from modules.import_files.model import ImportFile
from modules.purchase_orders.model import PurchaseOrder, POLineItem, PackingListItem
from modules.freight_booking.model import ShipmentBooking
from modules.cargo_shipping.model import CargoShippingRecord
from modules.import_companies.model import ImportCompany
from modules.suppliers.model import Supplier
from modules.import_documentation.model import (
    DraftBLReviewSession,
    CertificateOfOriginReviewSession,
    InspectionCertificateReviewSession,
)
from modules.import_documentation.schemas import (
    POFinalAdjustmentRequest,
    DraftBLReviewCreate,
    DraftBLReviewUpdate,
    DraftBLReviewResponse,
    DraftBLComparisonRequest,
    CertificateOfOriginReviewCreate,
    CertificateOfOriginReviewUpdate,
    CertificateOfOriginReviewResponse,
    COOComparisonRequest,
    InspectionCertificateReviewCreate,
    InspectionCertificateReviewUpdate,
    LegalDocsExpiryComplianceResponse,
    LegalDocAlertItem,
    InvoiceBLExtractAndMatchRequest,
    InvoiceBLExtractAndMatchResponse,
    InvoiceBLSyncRequest,
    POExtractAndCompareRequest,
    POExtractAndCompareResponse,
)




def _normalize_str(s: Any) -> str:
    if s is None:
        return ""
    st = str(s).strip()
    # Strip leading country-code + registration ID or numeric tax ID if present before company name
    st = re.sub(r'^[A-Z]{2}\d{6,15}[,\s]+|^\d{8,15}[,\s]+', '', st, flags=re.I)
    return " ".join(st.lower().split())


def _fuzzy_match(s1: Any, s2: Any, threshold: float = 0.70) -> tuple[bool, float]:
    n1 = _normalize_str(s1)
    n2 = _normalize_str(s2)
    if not n1 and not n2:
        return True, 1.0
    if not n1 or not n2:
        return False, 0.0
    if n1 == n2 or n1 in n2 or n2 in n1:
        return True, 1.0
    ratio = difflib.SequenceMatcher(None, n1, n2).ratio()
    w1 = set(n1.split())
    w2 = set(n2.split())
    if w1 and w2:
        overlap = len(w1.intersection(w2)) / max(len(w1), len(w2))
        if overlap >= 0.50:
            return True, max(ratio, overlap)
    return ratio >= threshold, ratio


def _numeric_match(val1: Any, val2: Any, tolerance_pct: float = 0.5) -> tuple[bool, float, float]:
    try:
        f1 = float(val1 or 0.0)
        f2 = float(val2 or 0.0)
    except (ValueError, TypeError):
        return False, 100.0, 0.0
    if f1 == 0.0 and f2 == 0.0:
        return True, 0.0, 0.0
    base = max(abs(f1), abs(f2))
    diff = abs(f1 - f2)
    pct = (diff / base) * 100.0 if base > 0 else 0.0
    return pct <= tolerance_pct, pct, diff


# --- 1. FINAL INVOICE & PACKING LIST PO RECONCILIATION ---
def reconcile_po_final_adjustments_service(db: Session, request: POFinalAdjustmentRequest) -> dict:
    imp_file = db.query(ImportFile).filter(ImportFile.import_file_id == request.import_file_id).first()
    if not imp_file:
        raise HTTPException(status_code=404, detail="Import File not found")

    updated_items = []
    total_net_weight = 0.0
    total_gross_weight = 0.0
    total_cbm = 0.0
    total_final_amount = 0.0

    for itm in request.items:
        final_price = itm.final_unit_price if itm.final_unit_price > 0 else itm.unit_price
        itm.unit_price = final_price
        itm.final_unit_price = final_price

        if itm.po_item_id:
            db_po_item = db.query(POLineItem).filter(POLineItem.item_id == itm.po_item_id).first()
            if db_po_item:
                db_po_item.quantity = itm.final_quantity
                if final_price > 0:
                    db_po_item.unit_price = final_price
                    db_po_item.total_price = itm.final_quantity * final_price
                if itm.final_net_weight_kg > 0:
                    db_po_item.net_weight_kg = itm.final_net_weight_kg
                if itm.final_gross_weight_kg > 0:
                    db_po_item.gross_weight_kg = itm.final_gross_weight_kg
                if itm.final_cbm > 0:
                    db_po_item.total_cbm = itm.final_cbm

            db_pl_item = db.query(PackingListItem).filter(PackingListItem.packing_item_id == itm.po_item_id).first()
            if db_pl_item:
                if itm.final_packages_count > 0:
                    db_pl_item.qty_pkg = itm.final_packages_count
                if itm.final_net_weight_kg > 0:
                    db_pl_item.total_net_weight_kg = itm.final_net_weight_kg
                if itm.final_gross_weight_kg > 0:
                    db_pl_item.total_gross_weight_kg = itm.final_gross_weight_kg
                if itm.final_cbm > 0:
                    db_pl_item.total_cbm = itm.final_cbm
                if itm.package_type:
                    db_pl_item.package_type = itm.package_type
        
        total_net_weight += itm.final_net_weight_kg
        total_gross_weight += itm.final_gross_weight_kg
        total_cbm += itm.final_cbm
        if final_price > 0:
            total_final_amount += itm.final_quantity * final_price
        
        qty_variance = 0.0
        if itm.initial_quantity > 0:
            qty_variance = ((itm.final_quantity - itm.initial_quantity) / itm.initial_quantity) * 100.0
        itm.variance_percentage = round(qty_variance, 2)

        price_variance = 0.0
        if itm.initial_unit_price > 0:
            price_variance = ((final_price - itm.initial_unit_price) / itm.initial_unit_price) * 100.0
        itm.price_variance_percentage = round(price_variance, 2)

        weight_variance = 0.0
        if itm.initial_gross_weight_kg > 0:
            weight_variance = ((itm.final_gross_weight_kg - itm.initial_gross_weight_kg) / itm.initial_gross_weight_kg) * 100.0
        itm.weight_variance_percentage = round(weight_variance, 2)

        updated_items.append(itm.model_dump())

    # Update Import File Snapshot
    imp_file.pi_number = request.final_invoice_number
    imp_file.total_amount = total_final_amount
    db.commit()

    return {
        "status": "success",
        "message": "Final Commercial Invoice & Packing List quantities, prices, and weights certified for downstream systems (B/L, Inventory in Transit, Warehouse Receiving).",
        "import_file_id": request.import_file_id,
        "final_invoice_number": request.final_invoice_number,
        "final_packing_list_number": request.final_packing_list_number,
        "total_items_count": len(updated_items),
        "total_net_weight_kg": round(total_net_weight, 2),
        "total_gross_weight_kg": round(total_gross_weight, 2),
        "total_cbm": round(total_cbm, 3),
        "total_final_amount": round(total_final_amount, 2),
        "items": updated_items,
    }


# --- 2. DRAFT BILL OF LADING (B/L) 5-STAGE REVIEW ENGINE ---
def _build_system_bl_snapshot(db: Session, import_file_id: int) -> dict:
    imp_file = db.query(ImportFile).filter(ImportFile.import_file_id == import_file_id).first()
    booking = db.query(ShipmentBooking).filter(ShipmentBooking.import_file_id == import_file_id, ShipmentBooking.is_active == True).first()
    cargo_shp = db.query(CargoShippingRecord).filter(CargoShippingRecord.import_file_id == import_file_id, CargoShippingRecord.is_active == True).first()
    company = db.query(ImportCompany).filter(ImportCompany.company_id == imp_file.company_id).first() if imp_file and imp_file.company_id else None
    supplier = db.query(Supplier).filter(Supplier.supplier_id == imp_file.supplier_id).first() if imp_file and imp_file.supplier_id else None
    pos = db.query(PurchaseOrder).filter(PurchaseOrder.import_file_id == import_file_id, PurchaseOrder.is_active == True).all()

    # Containers from Phase 5 Cargo Shipping
    containers = []
    cargo_gw = 0.0
    cargo_nw = 0.0
    hs_codes = set()
    goods_desc = []
    container_str_parts = []

    if cargo_shp and cargo_shp.containers_loading_data:
        for c in cargo_shp.containers_loading_data:
            c_dict = dict(c)
            c_no = c_dict.get("container_no", "")
            s_no = c_dict.get("seal_no", "")
            c_sz = c_dict.get("container_type", "40HC")
            containers.append({
                "container_no": c_no,
                "seal_no": s_no,
                "container_type": c_sz,
                "gross_weight_kg": float(c_dict.get("gross_weight_kg") or 0.0),
                "net_weight_kg": float(c_dict.get("net_weight_kg") or 0.0),
                "vgm_status": c_dict.get("vgm_status", "Submitted"),
            })
            cargo_gw += float(c_dict.get("gross_weight_kg") or 0.0)
            cargo_nw += float(c_dict.get("net_weight_kg") or 0.0)
            if c_no:
                container_str_parts.append(f"{c_no} / Seal: {s_no} ({c_sz})")

    has_any_pl = any(bool(po.packing_list_items) for po in pos)

    pl_total_cbm = 0.0
    pl_total_gw = 0.0
    pl_total_nw = 0.0
    pl_total_pkgs = 0
    pl_total_pcs = 0

    inv_total_cbm = 0.0
    inv_total_gw = 0.0
    inv_total_nw = 0.0
    inv_total_pkgs = 0
    inv_total_pcs = 0

    for po in pos:
        # 1. Line items
        if po.line_items:
            for itm in po.line_items:
                desc = itm.description_ar or itm.description_en or itm.item_code
                if desc and desc not in goods_desc:
                    goods_desc.append(desc)
                qty = float(itm.quantity or 0.0)
                inv_total_pcs += int(qty)
                inv_total_pkgs += int(qty)
                inv_total_cbm += float(itm.total_cbm or 0.0)
                inv_total_gw += float(itm.gross_weight_kg or 0.0)
                inv_total_nw += float(itm.net_weight_kg or 0.0)
                if itm.tariff and itm.tariff.hs_code:
                    hs_codes.add(itm.tariff.hs_code)

        # 2. Packing list items
        if po.packing_list_items:
            for pl in po.packing_list_items:
                p_code = pl.item_code or pl.hs_code
                if p_code and p_code not in goods_desc:
                    goods_desc.append(p_code)
                q_pkg = float(pl.qty_pkg or 0.0)
                q_pcs = float(pl.qty_pcs or 0.0)
                pl_total_pkgs += int(q_pkg)
                pl_total_pcs += int(q_pcs)

                # CBM
                cbm_val = float(pl.total_cbm or 0.0)
                if cbm_val <= 0:
                    l_m = float(pl.length_cm or 0.0) / 100.0
                    w_m = float(pl.width_cm or 0.0) / 100.0
                    h_m = float(pl.height_cm or 0.0) / 100.0
                    cbm_val = round(q_pkg * (l_m * w_m * h_m), 4) if (l_m > 0 and w_m > 0 and h_m > 0) else 0.0
                pl_total_cbm += cbm_val

                # Gross Weight
                gw_val = float(pl.total_gross_weight_kg or 0.0)
                if gw_val <= 0 and float(pl.gross_weight_unit_kg or 0.0) > 0:
                    gw_val = round(q_pkg * float(pl.gross_weight_unit_kg), 2)
                pl_total_gw += gw_val

                # Net Weight
                nw_val = float(pl.total_net_weight_kg or 0.0)
                if nw_val <= 0 and float(pl.net_weight_unit_kg or 0.0) > 0:
                    nw_val = round(q_pkg * float(pl.net_weight_unit_kg), 2)
                pl_total_nw += nw_val

                if pl.hs_code:
                    hs_codes.add(pl.hs_code)

    # Determine final totals with priority to detailed Packing List
    if has_any_pl and (pl_total_pkgs > 0 or pl_total_gw > 0 or pl_total_cbm > 0):
        total_cbm = pl_total_cbm
        total_gw = pl_total_gw
        total_nw = pl_total_nw
        total_pkgs = pl_total_pkgs
        total_pcs = pl_total_pcs if pl_total_pcs > 0 else inv_total_pcs
    else:
        total_cbm = inv_total_cbm
        total_gw = cargo_gw if cargo_gw > 0 else inv_total_gw
        total_nw = cargo_nw if cargo_nw > 0 else inv_total_nw
        total_pkgs = inv_total_pkgs
        total_pcs = inv_total_pcs

    shipper_name = supplier.company_name if supplier else (imp_file.supplier_name if imp_file else 'Foreign Supplier')
    shipper_addr = supplier.address if (supplier and supplier.address) else 'Export Industrial District, Global Port'
    shipper_phone = getattr(supplier, 'phone', None) or getattr(supplier, 'contact_phone', None) or '+49-89-636-00'
    shipper_full = f"{shipper_name}\nAddress: {shipper_addr}\nPhone: {shipper_phone}".strip()

    consignee_name = company.importer_name if company else (imp_file.company_name if imp_file else 'Importing Company')
    consignee_addr = company.address if (company and company.address) else '15 Industrial Zone, Cairo, Egypt'
    consignee_phone = getattr(company, 'phone', None) or '+20-2-25778899'
    consignee_full = f"{consignee_name}\nAddress: {consignee_addr}\nPhone: {consignee_phone}".strip()

    notify_party_full = consignee_full
    vessel = booking.vessel_name if (booking and booking.vessel_name) else "OCEAN VESSEL"
    voyage = getattr(booking, 'voyage_number', None) or getattr(booking, 'voyage_no', None) or "VOY-01"
    pol = booking.pol_name if (booking and booking.pol_name) else "Port of Loading (POL)"
    pod = booking.pod_name if (booking and booking.pod_name) else "Port of Discharge (POD)"
    freight_terms = getattr(booking, 'freight_terms', None) or "Freight Prepaid"
    place_of_deliv = getattr(booking, 'place_of_delivery', None) or pod
    bkg_no = booking.booking_confirmation_no if (booking and booking.booking_confirmation_no) else (cargo_shp.booking_no if (cargo_shp and cargo_shp.booking_no) else "BKG-REF")
    acid_no = imp_file.acid_number if (imp_file and imp_file.acid_number) else "N/A"
    tax_id = company.vat_id if company else "N/A"
    shipper_reg = getattr(supplier, 'registration_id', None) or getattr(supplier, 'tax_id', None) or "N/A"
    shp_mode = booking.shipment_type if (booking and booking.shipment_type) else "FCL / FCL"
    container_summary_str = "; ".join(container_str_parts) if container_str_parts else "N/A"

    return {
        "forwarder_name": booking.freight_forwarder_name if (booking and booking.freight_forwarder_name) else "Direct Shipping Line",
        "shipping_line": booking.shipping_line_name if (booking and booking.shipping_line_name) else (cargo_shp.shipping_line if (cargo_shp and cargo_shp.shipping_line) else "OCEAN CARRIER / FREIGHT LINE"),
        "vessel_name": vessel,
        "voyage_number": voyage,
        "pol": pol,
        "pod": pod,
        "freight_terms": freight_terms,
        "place_of_delivery": place_of_deliv,
        "booking_no": bkg_no,
        "acid_number": acid_no,
        "importer_tax_id": tax_id,
        "shipper_reg_id": shipper_reg,
        "shipping_mode": shp_mode,
        "container_summary": container_summary_str,
        "hbl_no": None,
        "mbl_no": None,
        "shipper": shipper_full,
        "consignee": consignee_full,
        "notify_party": notify_party_full,
        "hs_codes": list(hs_codes) if hs_codes else ["N/A"],
        "goods_description": ", ".join(goods_desc) if goods_desc else "General Merchandise & Import Goods",
        "qty_pcs": total_pcs,
        "qty_pkg": total_pkgs,
        "total_net_weight_kg": round(total_nw, 2),
        "total_gross_weight_kg": round(total_gw, 2),
        "cbm": round(total_cbm, 4),
        "containers": containers,
        "container_count": len(containers),
    }


def _ocr_pdf_or_image(filename: str, content_bytes: bytes) -> str:
    """
    High-accuracy optical character recognition (OCR) for scanned PDFs and images.
    Uses RapidOCR + pypdfium2 / Pillow.
    """
    lower = filename.lower()
    try:
        from rapidocr_onnxruntime import RapidOCR
        import numpy as np
        engine = RapidOCR()
        extracted_lines = []

        if lower.endswith(".pdf"):
            import pypdfium2 as pdfium
            pdf = pdfium.PdfDocument(content_bytes)
            for i in range(len(pdf)):
                page = pdf.get_page(i)
                pil_img = page.render(scale=2.0).to_pil()
                ocr_res, _ = engine(np.array(pil_img))
                if ocr_res:
                    extracted_lines.extend([item[1] for item in ocr_res])
        elif lower.endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp")):
            from PIL import Image
            pil_img = Image.open(io.BytesIO(content_bytes)).convert("RGB")
            ocr_res, _ = engine(np.array(pil_img))
            if ocr_res:
                extracted_lines.extend([item[1] for item in ocr_res])

        return "\n".join(extracted_lines)
    except Exception as e:
        logger.warning(f"RapidOCR extraction failed for '{filename}': {e}")
        return ""


def extract_text_and_boxes_from_uploaded_file(filename: str, content_bytes: bytes) -> Tuple[str, dict]:
    """
    Extracts raw text and 2D spatial bounding boxes from uploaded PDF or files.
    Includes automatic OCR fallback for scanned documents and images.
    """
    lower_name = filename.lower()
    text_content = ""
    spatial_boxes: dict = {}

    if lower_name.endswith(".pdf"):
        try:
            from modules.import_documentation.ai_document_parser import extract_spatial_pdf_text_and_boxes
            text_content, spatial_boxes = extract_spatial_pdf_text_and_boxes(content_bytes)
        except Exception as e:
            logger.warning(f"Spatial PDF extraction failed: {e}. Falling back to pypdf.")
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(content_bytes))
                parts = []
                for page in reader.pages:
                    try:
                        t = page.extract_text(extraction_mode="layout")
                    except Exception:
                        t = page.extract_text()
                    if t:
                        parts.append(t)
                text_content = "\n\n".join(parts)
            except Exception as e2:
                text_content = f"PDF Extraction Note: {str(e2)}"
    else:
        text_content = extract_text_from_uploaded_file(filename, content_bytes)

    # Automatic OCR fallback if document contains scanned images or no text stream
    if not text_content or len(text_content.strip()) < 30:
        ocr_text = _ocr_pdf_or_image(filename, content_bytes)
        if ocr_text and len(ocr_text.strip()) > len(text_content.strip()):
            text_content = ocr_text

    return text_content, spatial_boxes


def extract_text_from_uploaded_file(filename: str, content_bytes: bytes) -> str:
    """
    Extracts raw text and table structures from uploaded PDF, Word (.docx), Excel (.xlsx), or Text files.
    """
    lower_name = filename.lower()
    text_content = ""

    if lower_name.endswith(".pdf"):
        try:
            from modules.import_documentation.ai_document_parser import extract_spatial_pdf_text_and_boxes
            spatial_text, _ = extract_spatial_pdf_text_and_boxes(content_bytes)
            text_content = spatial_text
        except Exception as e:
            logger.warning(f"Spatial PDF extraction failed: {e}. Falling back to pypdf.")
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(content_bytes))
                parts = []
                for page in reader.pages:
                    try:
                        t = page.extract_text(extraction_mode="layout")
                    except Exception:
                        t = page.extract_text()
                    if t:
                        parts.append(t)
                text_content = "\n\n".join(parts)
            except Exception as e2:
                text_content = f"PDF Extraction Note: {str(e2)}"

    elif lower_name.endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp")):
        text_content = _ocr_pdf_or_image(filename, content_bytes)

    elif lower_name.endswith(".docx") or lower_name.endswith(".doc"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content_bytes))
            parts = [p.text for p in doc.paragraphs if p.text]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()]))
            text_content = "\n".join(parts)
        except Exception as e:
            text_content = f"Word Document Extraction Note: {str(e)}"

    elif lower_name.endswith(".xlsx") or lower_name.endswith(".xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content_bytes), data_only=True)
            parts = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows(values_only=True):
                    row_vals = [str(v).strip() for v in row if v is not None and str(v).strip() != ""]
                    if row_vals:
                        parts.append(" | ".join(row_vals))
            text_content = "\n".join(parts)
        except Exception as e:
            text_content = f"Excel Extraction Note: {str(e)}"

    else:
        try:
            text_content = content_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text_content = content_bytes.decode("latin-1", errors="ignore")

    return text_content


def parse_draft_bl_raw_text(raw_text: str, spatial_boxes: Optional[dict] = None) -> dict:
    from modules.import_documentation.ai_document_parser import extract_draft_bl_with_ai
    return extract_draft_bl_with_ai(raw_text, spatial_boxes=spatial_boxes)




def compare_draft_bl_service(db: Session, request: DraftBLComparisonRequest) -> dict:
    sys_data = _build_system_bl_snapshot(db, request.import_file_id)
    draft_fields = dict(request.draft_fields or {})

    raw_input = getattr(request, "raw_draft_text", None) or getattr(request, "raw_text", None)
    if raw_input:
        extracted = parse_draft_bl_raw_text(raw_input)
        for k, v in extracted.items():
            if k not in draft_fields or not draft_fields[k]:
                draft_fields[k] = v

    if "container_summary" not in draft_fields and "containers" in draft_fields:
        c_parts = []
        for c in draft_fields["containers"]:
            c_no = c.get("container_no", "")
            s_no = c.get("seal_no", "")
            if c_no:
                c_parts.append(f"{c_no} / Seal: {s_no}" if s_no else c_no)
        if c_parts:
            draft_fields["container_summary"] = "; ".join(c_parts)

    # Field Mappings & Rules
    fields_spec = [
        ("shipper", "المصدر / الشاحن (Shipper)", "Shipper Name, Address & Phone", "Supplier Master Data", False, "fuzzy", "Supplier"),
        ("consignee", "المستورد / المرسل إليه (Consignee)", "Consignee Name, Address & Phone", "Importer Company Master Data", True, "fuzzy", "Importer"),
        ("notify_party", "جهة الإخطار (Notify Party)", "Notify Party Name, Address & Phone", "Importer / Contract", False, "fuzzy", "Importer"),
        ("vessel_name", "اسم الباخرة (Vessel)", "Vessel Name", "Final Booking", False, "text", "Shipping Provider"),
        ("voyage_number", "رقم الرحلة (Voyage)", "Voyage Number", "Final Booking", False, "text", "Shipping Provider"),
        ("pol", "ميناء الشحن (POL)", "Port of Loading (POL)", "Final Booking", False, "text", "Shipping Provider"),
        ("pod", "ميناء التفريغ (POD)", "Port of Discharge (POD)", "Final Booking", False, "text", "Shipping Provider"),
        ("freight_terms", "شروط النولون (Freight Terms)", "Freight Terms (Prepaid/Collect)", "Final Booking", False, "text", "Shipping Provider"),
        ("place_of_delivery", "مكان التسليم (Place of Delivery)", "Place of Delivery", "Final Booking", False, "text", "Shipping Provider"),
        ("booking_no", "رقم الحجز (Booking Number)", "Booking Number", "Final Booking", True, "text", "Shipping Provider"),
        ("acid_number", "رقم القيد الجمركي المبدئي (ACID)", "ACID Number", "Import File & Nafeza", True, "text", "Importer"),
        ("importer_tax_id", "البطاقة الضريبية للمستورد (Importer Tax ID)", "Egyptian Importer Tax ID", "Importer Company", True, "text", "Importer"),
        ("shipper_reg_id", "رقم تسجيل المصدر (Shipper Reg ID)", "Shipper Registration/ID", "Supplier", False, "text", "Supplier"),
        ("shipping_mode", "نمط الشحن (Shipping Mode)", "Shipping Mode (FCL/LCL/Air)", "Final Booking", False, "text", "Shipping Provider"),
        ("container_summary", "بيان الحاويات والسيل (Container Summary)", "Container No + Seal No + Size", "Container Allocation", True, "fuzzy", "Shipping Provider"),
        ("goods_description", "بيان التعبئة والوصف (Packing Summary)", "Packing Summary & Goods Description", "Final Approved Packing List", False, "fuzzy", "Supplier"),
        ("total_gross_weight_kg", "الوزن القائم (Gross Weight)", "Gross Weight (KG)", "Final Approved Packing List", False, "numeric_05", "Shipping Provider"),
        ("total_net_weight_kg", "الوزن الصافي (Net Weight)", "Net Weight (KG)", "Final Approved Packing List", False, "numeric_05", "Supplier"),
        ("cbm", "الحجم الإجمالي (Measurement CBM)", "Measurement (CBM)", "Final Approved Packing List", False, "numeric_05", "Shipping Provider"),
        ("qty_pkg", "عدد الطرود (Number of Packages)", "Number of Packages", "Final Approved Packing List", False, "numeric_strict", "Supplier"),
    ]

    matrix = []
    checklist = []
    blocking_reasons = []

    for key, label_ar, label_en, source, is_critical, match_type, resp_default in fields_spec:
        sys_val = sys_data.get(key)
        draft_val = draft_fields.get(key)

        if draft_val is None or str(draft_val).strip() == "":
            match_status = "MISSING_IN_DRAFT"
            is_explicitly_empty = key in draft_fields and draft_fields[key] == ""
            severity = "BLOCKING" if (is_critical and is_explicitly_empty) else "WARNING"
            details = "الحقل مطلوب ولم يتم استخراجه أو إدخاله في بيانات درافت البوليصة."
            if is_critical and is_explicitly_empty:
                blocking_reasons.append(f"حقل حرج فارغ: {label_ar}")
            tolerance_pct = 0.0
        elif sys_val is None or str(sys_val).strip() == "":
            match_status = "MATCH"
            severity = "NONE"
            details = f"تم استخراج القيمة من الدرافت ({draft_val})."
            tolerance_pct = 0.0
        elif match_type == "fuzzy":
            matched, ratio = _fuzzy_match(sys_val, draft_val, threshold=0.85)
            tolerance_pct = round((1.0 - ratio) * 100.0, 1)
            if matched:
                match_status = "MATCH"
                severity = "NONE"
                details = f"مطابقة ذكية ممتازة بنسبة {round(ratio * 100, 1)}%."
            else:
                match_status = "MISMATCH_CRITICAL" if is_critical else "MISMATCH_MINOR"
                severity = "BLOCKING" if is_critical else "WARNING"
                details = f"عدم تطابق! قيمة النظام: '{sys_val}' | قيمة الدرافت: '{draft_val}'."
                if is_critical:
                    blocking_reasons.append(f"اختلاف حرج في {label_ar}")
        elif match_type == "numeric_05":
            try:
                s_num = float(sys_val)
                d_num = float(draft_val)
                diff = abs(s_num - d_num)
                pct = (diff / s_num * 100.0) if s_num > 0 else 0.0
                tolerance_pct = round(pct, 2)
                if pct <= 0.5:
                    match_status = "MATCH"
                    severity = "NONE"
                    details = f"مطابقة رقمية ممتازة (فرق {tolerance_pct}% ضمن نسبة السماح 0.5%)."
                else:
                    match_status = "MISMATCH_MINOR"
                    severity = "WARNING"
                    details = f"فرق أوزان يتجاوز نسبة السماح ({tolerance_pct}%). النظام: {s_num} | الدرافت: {d_num}."
            except Exception:
                match_status = "MISMATCH_MINOR"
                severity = "WARNING"
                details = f"تعذر التحقق الرقمي من القيمة: {draft_val}"
        else: # text
            s_clean = str(sys_val).strip().upper()
            d_clean = str(draft_val).strip().upper()
            if s_clean == d_clean:
                match_status = "MATCH"
                severity = "NONE"
                details = "مطابقة نصية تامة 100%."
            else:
                match_status = "MISMATCH_CRITICAL" if is_critical else "MISMATCH_MINOR"
                severity = "BLOCKING" if is_critical else "WARNING"
                details = f"عدم تطابق! القيمة المعتمدة: '{sys_val}' | قيمة المسودة: '{draft_val}'."
                if is_critical:
                    blocking_reasons.append(f"اختلاف حرج في {label_ar}")

        matrix.append({
            "field_key": key,
            "field_label_ar": label_ar,
            "field_label_en": label_en,
            "source_entity": source,
            "system_value": sys_val,
            "draft_value": draft_val,
            "match_status": match_status,
            "severity": severity,
            "tolerance_pct": tolerance_pct,
            "details": details,
        })

        is_match = (match_status == "MATCH")
        checklist.append({
            "field_key": key,
            "field_label_ar": label_ar,
            "field_label_en": label_en,
            "source_entity": source,
            "system_value": sys_val,
            "draft_value": draft_val,
            "status": "Correct" if is_match else "Incorrect",
            "required_correction": f"Correct {label_en} to: {sys_val}" if not is_match else None,
            "reason": details if not is_match else None,
            "notes": None,
            "responsible_party": resp_default,
            "is_locked": False,
            "previous_status": None,
        })

    has_blocking = len(blocking_reasons) > 0
    open_discrepancies = [c for c in checklist if c["status"] == "Incorrect"]

    # Generate Revision Report Items
    revision_report = []
    for itm in open_discrepancies:
        revision_report.append({
            "item": itm["field_label_en"],
            "required_action": itm["required_correction"] or f"Update {itm['field_label_en']} to match system value ({itm['system_value']})",
            "responsible": itm["responsible_party"] or "Shipping Provider",
            "reason": itm["reason"] or "Value mismatch with master reference",
            "notes": itm["notes"],
        })

    # Auto-generate Correction Letter
    correction_letter = ""
    if revision_report:
        correction_letter = f"""Subject: Urgent: Draft B/L Correction Request - Booking Ref: {sys_data.get('booking_no')} / ACID: {sys_data.get('acid_number')}

Dear Shipping Line / Forwarder Operations Team,

Please be advised that upon verification of the Draft Bill of Lading received for Booking No: {sys_data.get('booking_no')}, the following discrepancies were identified and require immediate amendment prior to final B/L issuance:

"""
        for idx, item in enumerate(revision_report, 1):
            correction_letter += f"{idx}. Item: {item['item']}\n"
            correction_letter += f"   - Required Action: {item['required_action']}\n"
            correction_letter += f"   - Responsible Party: {item['responsible']}\n"
            correction_letter += f"   - Reason / Details: {item['reason']}\n\n"
        
        correction_letter += "Please provide the revised Draft B/L with the requested corrections at your earliest convenience.\n\nBest Regards,\nImport Operations Department"

    stage_calc = "Stage 2: Revision Required" if open_discrepancies else "Stage 4: Dual Approval"
    status_calc = "REVISION_REQUIRED" if open_discrepancies else "REVIEWED_PENDING_APPROVAL"

    return {
        "import_file_id": request.import_file_id,
        "draft_source": request.draft_source,
        "stage": stage_calc,
        "system_snapshot_data": sys_data,
        "draft_input_data": draft_fields,
        "comparison_matrix": matrix,
        "checklist_data": checklist,
        "revision_report_data": revision_report,
        "has_blocking_mismatch": has_blocking,
        "open_discrepancies_count": len(open_discrepancies),
        "blocking_reasons": blocking_reasons,
        "correction_request_letter": correction_letter,
        "status": status_calc,
    }


def create_draft_bl_review_service(db: Session, schema: DraftBLReviewCreate) -> DraftBLReviewSession:
    if schema.checklist_data:
        # Convert any Pydantic models to dicts
        raw_checklist = []
        revision_report = []
        open_count = 0
        for item in schema.checklist_data:
            item_dict = item.model_dump() if hasattr(item, "model_dump") else (item.dict() if hasattr(item, "dict") else item)
            status_val = item_dict.get("status", "Correct")
            if status_val == "Incorrect":
                open_count += 1
                revision_report.append({
                    "item": item_dict.get("field_label_ar") or item_dict.get("field_label_en", ""),
                    "required_action": item_dict.get("required_correction", ""),
                    "responsible": item_dict.get("responsible_party") or "Shipping Provider",
                    "reason": item_dict.get("reason", ""),
                    "notes": item_dict.get("notes"),
                })
            raw_checklist.append(item_dict)
        schema.checklist_data = raw_checklist
        schema.revision_report_data = revision_report
        schema.open_discrepancies_count = open_count
        if open_count > 0:
            schema.stage = "Stage 2: Revision Required"
            schema.status = "REVISION_REQUIRED"
        else:
            schema.stage = "Stage 4: Dual Approval"
            schema.status = "REVIEWED_PENDING_APPROVAL"
    elif not schema.comparison_matrix:
        comp_req = DraftBLComparisonRequest(
            import_file_id=schema.import_file_id,
            draft_source=schema.draft_source,
            draft_fields=schema.draft_input_data,
        )
        comp_res = compare_draft_bl_service(db, comp_req)
        schema.system_snapshot_data = comp_res["system_snapshot_data"]
        schema.comparison_matrix = comp_res["comparison_matrix"]
        schema.checklist_data = comp_res["checklist_data"]
        schema.revision_report_data = comp_res["revision_report_data"]
        schema.has_blocking_mismatch = comp_res["has_blocking_mismatch"]
        schema.open_discrepancies_count = comp_res["open_discrepancies_count"]
        schema.blocking_reasons = comp_res["blocking_reasons"]
        schema.correction_request_letter = comp_res["correction_request_letter"]
        schema.stage = comp_res["stage"]
        schema.status = comp_res["status"]

    existing = repo.get_draft_bl_review_by_file_id(db, schema.import_file_id, include_inactive=False)
    if existing:
        update_schema = DraftBLReviewUpdate(**schema.model_dump(exclude_unset=True))
        return repo.update_draft_bl_review(db, existing.bl_review_id, update_schema)
    return repo.create_draft_bl_review(db, schema)


def update_draft_bl_review_service(db: Session, review_id: int, schema: DraftBLReviewUpdate) -> DraftBLReviewSession:
    return repo.update_draft_bl_review(db, review_id, schema)


def update_draft_bl_checklist_service(db: Session, review_id: int, checklist_items: List[DraftBLChecklistItem], reviewer_name: str = "Kamal") -> DraftBLReviewSession:
    review = repo.get_draft_bl_review_by_id(db, review_id)
    if not review:
        raise HTTPException(status_code=404, detail="Draft B/L Review not found")

    updated_checklist = []
    revision_report = []
    open_count = 0

    for item in checklist_items:
        item_dict = item.model_dump()
        if item.status == "Incorrect":
            if not item.required_correction or not item.reason:
                raise HTTPException(
                    status_code=400,
                    detail=f"الحقول 'Required Correction' و 'Reason' إلزامية للبند غير المطابق: {item.field_label_ar}"
                )
            open_count += 1
            revision_report.append({
                "item": item.field_label_en,
                "required_action": item.required_correction,
                "responsible": item.responsible_party or "Shipping Provider",
                "reason": item.reason,
                "notes": item.notes,
            })
        updated_checklist.append(item_dict)

    review.checklist_data = updated_checklist
    review.revision_report_data = revision_report
    review.open_discrepancies_count = open_count
    review.reviewed_by = reviewer_name
    review.reviewed_at = datetime.now(timezone.utc)

    # Re-generate correction letter
    if revision_report:
        sys_data = review.system_snapshot_data or {}
        correction_letter = f"""Subject: Urgent: Draft B/L Correction Request - Booking Ref: {sys_data.get('booking_no')} / ACID: {sys_data.get('acid_number')}

Dear Shipping Line / Forwarder Operations Team,

Please be advised that upon review of the Draft Bill of Lading received for Booking No: {sys_data.get('booking_no')}, the following corrections are required:

"""
        for idx, itm in enumerate(revision_report, 1):
            correction_letter += f"{idx}. {itm['item']}:\n"
            correction_letter += f"   - Required Action: {itm['required_action']}\n"
            correction_letter += f"   - Responsible Party: {itm['responsible']}\n"
            correction_letter += f"   - Reason: {itm['reason']}\n\n"
        correction_letter += "Please reissue the draft B/L with these corrections at your earliest convenience.\n\nBest Regards,\nImport Operations Department"
        review.correction_request_letter = correction_letter
        review.stage = "Stage 2: Revision Required"
        review.status = "REVISION_REQUIRED"
    else:
        review.stage = "Stage 4: Dual Approval"
        review.status = "REVIEWED_PENDING_APPROVAL"

    db.commit()
    db.refresh(review)
    return review


def process_dual_approval_service(db: Session, request: DualApprovalRequest) -> DraftBLReviewSession:
    review = repo.get_draft_bl_review_by_id(db, request.bl_review_id)
    if not review:
        raise HTTPException(status_code=404, detail="Draft B/L Review not found")

    if review.open_discrepancies_count > 0:
        raise HTTPException(
            status_code=400,
            detail="لا يمكن اعتماد درافت البوليصة طالما توجد بنود غير مطابقة (Open Discrepancies) لم يتم إغلاقها."
        )

    now = datetime.now(timezone.utc)

    if request.role == "importer":
        if request.action == "Approved":
            review.importer_approval_status = "Approved"
            review.importer_approved_by = request.approved_by
            review.importer_approval_date = now
            review.importer_approval_notes = request.notes
        else:
            review.importer_approval_status = "Rejected"
            review.importer_approval_notes = request.notes
            review.stage = "Stage 2: Revision Required"
            review.status = "REVISION_REQUIRED"
    elif request.role == "customs_broker":
        if request.action == "Approved":
            review.broker_approval_status = "Approved"
            review.broker_approved_by = request.approved_by
            review.broker_approval_date = now
            review.broker_approval_notes = request.notes
        else:
            review.broker_approval_status = "Rejected"
            review.broker_approval_notes = request.notes
            review.stage = "Stage 2: Revision Required"
            review.status = "REVISION_REQUIRED"
    else:
        raise HTTPException(status_code=400, detail="Invalid approval role. Must be 'importer' or 'customs_broker'.")

    # If BOTH approved -> Stage 5 Final
    if review.importer_approval_status == "Approved" and review.broker_approval_status == "Approved":
        review.stage = "Stage 5: Final"
        review.status = "FINAL"
        review.approved_by = f"{review.importer_approved_by} & {review.broker_approved_by}"
        review.approved_at = now

    db.commit()
    db.refresh(review)
    return review


def create_new_draft_version_service(db: Session, request: NewDraftVersionRequest) -> DraftBLReviewSession:
    parent = repo.get_draft_bl_review_by_id(db, request.parent_session_id)
    if not parent:
        raise HTTPException(status_code=404, detail="Parent Draft B/L Review not found")

    # Run new comparison
    comp_req = DraftBLComparisonRequest(
        import_file_id=parent.import_file_id,
        draft_source=request.draft_source,
        raw_draft_text=request.raw_draft_text,
        draft_fields=request.draft_fields,
    )
    comp_res = compare_draft_bl_service(db, comp_req)

    # Carry forward locks for previously approved items whose draft value didn't change
    parent_checklist_map = {c["field_key"]: c for c in (parent.checklist_data or [])}
    new_checklist = []

    for itm in comp_res["checklist_data"]:
        f_key = itm["field_key"]
        prev_itm = parent_checklist_map.get(f_key)
        if prev_itm and prev_itm.get("status") == "Correct" and str(prev_itm.get("draft_value")) == str(itm.get("draft_value")):
            itm["status"] = "Correct"
            itm["is_locked"] = True
            itm["previous_status"] = "Correct"
        else:
            itm["is_locked"] = False
            itm["previous_status"] = prev_itm.get("status") if prev_itm else None
        new_checklist.append(itm)

    new_version_num = parent.version_number + 1
    new_code = f"{parent.bl_review_code}-v{new_version_num}"

    new_session = DraftBLReviewSession(
        bl_review_code=new_code,
        import_file_id=parent.import_file_id,
        po_id=parent.po_id,
        booking_id=parent.booking_id,
        draft_bl_number=f"{parent.draft_bl_number}-v{new_version_num}",
        shipping_line=parent.shipping_line,
        vessel_name=parent.vessel_name,
        voyage_number=parent.voyage_number,
        booking_no=parent.booking_no,
        hbl_no=parent.hbl_no,
        mbl_no=parent.mbl_no,
        freight_terms=parent.freight_terms,
        place_of_delivery=parent.place_of_delivery,
        importer_tax_id=parent.importer_tax_id,
        shipper_reg_id=parent.shipper_reg_id,
        measurement_cbm=parent.measurement_cbm,
        net_weight_kg=parent.net_weight_kg,
        packages_count=parent.packages_count,
        container_summary=parent.container_summary,
        draft_source=request.draft_source,
        version_number=new_version_num,
        parent_session_id=parent.bl_review_id,
        stage="Stage 3: Reviewed",
        system_snapshot_data=comp_res["system_snapshot_data"],
        draft_input_data=request.draft_fields,
        comparison_matrix=comp_res["comparison_matrix"],
        checklist_data=new_checklist,
        revision_report_data=comp_res["revision_report_data"],
        has_blocking_mismatch=comp_res["has_blocking_mismatch"],
        open_discrepancies_count=len([c for c in new_checklist if c["status"] == "Incorrect"]),
        blocking_reasons=comp_res["blocking_reasons"],
        correction_request_letter=comp_res["correction_request_letter"],
        status="REVIEWED_PENDING_APPROVAL" if len([c for c in new_checklist if c["status"] == "Incorrect"]) == 0 else "REVISION_REQUIRED",
    )
    db.add(new_session)
    db.commit()
    db.refresh(new_session)
    return new_session


def approve_draft_bl_service(db: Session, review_id: int, approved_by: str = "Kamal") -> DraftBLReviewSession:
    review = repo.get_draft_bl_review_by_id(db, review_id)
    if not review:
        raise HTTPException(status_code=404, detail="Draft B/L Review not found")
    if review.open_discrepancies_count > 0 or review.has_blocking_mismatch:
        raise HTTPException(
            status_code=400,
            detail=f"لا يمكن اعتماد درافت البوليصة لوجود بنود غير مطابقة أو أخطاء حرجة (Blocking Discrepancies)."
        )
    review.stage = "Stage 5: Final"
    review.status = "FINAL"
    review.importer_approval_status = "Approved"
    review.importer_approved_by = approved_by
    review.importer_approval_date = datetime.now(timezone.utc)
    review.broker_approval_status = "Approved"
    review.broker_approved_by = approved_by
    review.broker_approval_date = datetime.now(timezone.utc)
    review.approved_by = approved_by
    review.approved_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(review)
    return review


def _compare_origins(sys_origin: str, drf_origin: str) -> Tuple[bool, float]:
    s_norm = sys_origin.lower().strip()
    d_norm = drf_origin.lower().strip()
    if not s_norm and not d_norm:
        return True, 1.0
    if not s_norm or not d_norm:
        return False, 0.0
    if s_norm == d_norm or s_norm in d_norm or d_norm in s_norm:
        return True, 1.0
    if ("eu" in d_norm or "european union" in d_norm or "الاتحاد الأوروبي" in d_norm) and any(eu == s_norm or eu in s_norm for eu in EU_27_COUNTRIES):
        return True, 1.0
    return _fuzzy_match(sys_origin, drf_origin, threshold=0.75)


# --- 3. CERTIFICATE OF ORIGIN (COO / EUR.1) VERIFICATION ENGINE ---
def compare_coo_service(db: Session, request: COOComparisonRequest) -> dict:
    imp_file = db.query(ImportFile).filter(ImportFile.import_file_id == request.import_file_id).first()
    if not imp_file:
        raise HTTPException(status_code=404, detail="Import File not found")
    company = db.query(ImportCompany).filter(ImportCompany.company_id == imp_file.company_id).first() if imp_file.company_id else None
    supplier = db.query(Supplier).filter(Supplier.supplier_id == imp_file.supplier_id).first() if imp_file.supplier_id else None
    raw_country = (getattr(supplier, 'foreign_exporter_country', None)) or (getattr(imp_file, 'origin_country', None)) or "Lithuania"
    country_iso = getattr(supplier, 'foreign_exporter_country_code', None) or COUNTRY_ISO_MAP.get(str(raw_country).strip().lower(), "")
    raw_reg = (getattr(supplier, 'tax_id', None)) or (getattr(supplier, 'foreign_exporter_id', None)) or (getattr(supplier, 'exporter_code', None)) or ""
    if raw_reg and country_iso and not str(raw_reg).strip().upper().startswith(country_iso.upper()):
        exp_reg_id = f"{country_iso.upper()}{str(raw_reg).strip()}"
    else:
        exp_reg_id = str(raw_reg).strip()

    sys_snapshot = {
        "exporter_name": supplier.company_name if supplier else imp_file.supplier_name,
        "exporter_reg_id": exp_reg_id,
        "exporter_address": supplier.address if supplier else "",
        "country_of_origin": raw_country,
        "importer_name": company.importer_name if company else imp_file.company_name,
        "importer_address": company.address if company else "",
        "destination_country": "Egypt",
        "invoice_number": imp_file.pi_number or "INV-2026-FINAL",
        "invoice_date": str(date.today()),
        "certificate_type": request.certificate_type,
    }

    draft_fields = request.draft_fields or {}
    matrix = []
    has_critical = False
    has_discrepancies = False

    coo_checks = [
        ("exporter_reg_id", "كود المصدر الأجنبي / رقم التسجيل (Foreign Exporter Reg Code)", False),
        ("exporter_name", "اسم المصدر (Exporter Name)", True),
        ("country_of_origin", "بلد المنشأ (Country of Origin)", True),
        ("importer_name", "اسم المستورد (Importer / Consignee)", True),
        ("destination_country", "بلد المقصد (Destination Country)", True),
        ("invoice_number", "رقم الفاتورة (Commercial Invoice No)", True),
        ("certificate_type", "نوع شهادة المنشأ (Certificate Type)", False),
    ]

    from modules.import_documentation.ai_document_parser import clean_exporter_name, clean_consignee_name
    cert_type_str = str(request.certificate_type or "").upper()
    is_china_or_standard_cert = bool("CHINA" in cert_type_str or "CCPIT" in cert_type_str or "STANDARD" in cert_type_str)

    for key, label_ar, is_critical in coo_checks:
        sys_val = sys_snapshot.get(key)
        drf_val = draft_fields.get(key)

        clean_sys_val = sys_val
        clean_drf_val = drf_val

        if key == "exporter_name":
            clean_sys_val = clean_exporter_name(str(sys_val or ""))
            clean_drf_val = clean_exporter_name(str(drf_val or ""), str(draft_fields.get("exporter_reg_id") or ""))
        elif key == "importer_name":
            clean_sys_val = clean_consignee_name(str(sys_val or ""))
            clean_drf_val = clean_consignee_name(str(drf_val or ""))

        if key == "exporter_reg_id":
            if is_china_or_standard_cert or not drf_val or not str(drf_val).strip():
                matched = True
                ratio = 1.0
            else:
                matched, ratio = _fuzzy_match(clean_sys_val, clean_drf_val, threshold=0.85)
        else:
            matched, ratio = _fuzzy_match(clean_sys_val, clean_drf_val, threshold=0.85)
        
        # Special rule for EUR.1 origin comparison: 'EU' or 'European Union' matches any EU-27 member country in system
        if key == "country_of_origin" and not matched:
            sys_norm = str(sys_val or "").lower().strip()
            drf_norm = str(drf_val or "").lower().strip()
            if ("eu" in drf_norm or "european union" in drf_norm or "الاتحاد الأوروبي" in drf_norm) and any(eu == sys_norm or eu in sys_norm for eu in EU_27_COUNTRIES):
                matched = True
                ratio = 1.0
        
        if matched:
            match_status = "MATCH"
            severity = "NONE"
        else:
            has_discrepancies = True
            if is_critical:
                has_critical = True
                match_status = "MISMATCH_CRITICAL"
                severity = "BLOCKING"
            else:
                match_status = "MISMATCH_MINOR"
                severity = "WARNING"

        displayed_draft_val = drf_val
        if key == "exporter_name":
            displayed_draft_val = clean_drf_val
        elif key == "importer_name":
            displayed_draft_val = clean_drf_val

        matrix.append({
            "field_key": key,
            "field_label_ar": label_ar,
            "system_value": sys_val,
            "draft_value": displayed_draft_val,
            "match_status": match_status,
            "severity": severity,
            "details": f"نسبة التشابه: {round(ratio * 100, 1)}%",
        })

    status_calc = "Verified" if not has_discrepancies else ("Discrepancy_Accepted" if not has_critical else "Correction Requested")

    return {
        "import_file_id": request.import_file_id,
        "certificate_type": request.certificate_type,
        "system_snapshot_data": sys_snapshot,
        "draft_input_data": draft_fields,
        "comparison_matrix": matrix,
        "has_discrepancies": has_discrepancies,
        "has_critical_mismatch": has_critical,
        "status": status_calc,
    }


def create_coo_review_service(db: Session, schema: CertificateOfOriginReviewCreate) -> CertificateOfOriginReviewSession:
    if not schema.comparison_matrix:
        comp = compare_coo_service(db, COOComparisonRequest(
            import_file_id=schema.import_file_id,
            certificate_type=schema.certificate_type,
            draft_fields=schema.draft_input_data,
        ))
        schema.system_snapshot_data = comp["system_snapshot_data"]
        schema.comparison_matrix = comp["comparison_matrix"]
        schema.has_discrepancies = comp["has_discrepancies"]
        schema.has_critical_mismatch = comp["has_critical_mismatch"]
        schema.status = comp["status"]

    # Auto-populate exporter, importer, country of origin if missing
    draft_dict = schema.draft_input_data or {}
    sys_dict = schema.system_snapshot_data or {}
    if not schema.exporter_name or schema.exporter_name == "N/A":
        schema.exporter_name = draft_dict.get("exporter_name") or sys_dict.get("exporter_name") or "Exporter"
    if not schema.importer_name or schema.importer_name == "N/A":
        schema.importer_name = draft_dict.get("importer_name") or sys_dict.get("importer_name") or "Importer"
    if not schema.country_of_origin or schema.country_of_origin == "N/A":
        schema.country_of_origin = draft_dict.get("country_of_origin") or sys_dict.get("country_of_origin") or "China"

    # Mandatory justification validation on discrepancies
    if (schema.has_discrepancies or schema.has_critical_mismatch):
        if schema.status in ["Verified", "Approved", "Discrepancy_Accepted"] and not (schema.override_reason and schema.override_reason.strip()):
            raise HTTPException(
                status_code=400,
                detail="يجب ذكر سبب ومبررات الموافقة على الاختلافات قبل اعتماد وحفظ دراسة شهادة المنشأ، أو العودة للتعديل ومخاطبة المورد."
            )

    existing = repo.get_coo_review_by_file_id(db, schema.import_file_id, include_inactive=False)
    if existing:
        update_schema = CertificateOfOriginReviewUpdate(**schema.model_dump(exclude_unset=True))
        return repo.update_coo_review(db, existing.coo_review_id, update_schema)
    return repo.create_coo_review(db, schema)


def update_coo_review_service(db: Session, review_id: int, schema: CertificateOfOriginReviewUpdate) -> CertificateOfOriginReviewSession:
    return repo.update_coo_review(db, review_id, schema)


# --- 4. INSPECTION CERTIFICATE (COC / COA / VOC) VERIFICATION ENGINE ---
def compare_inspection_cert_service(db: Session, request: InspectionComparisonRequest) -> dict:
    imp_file = db.query(ImportFile).filter(ImportFile.import_file_id == request.import_file_id).first()
    if not imp_file:
        raise HTTPException(status_code=404, detail="Import File not found")
    company = db.query(ImportCompany).filter(ImportCompany.company_id == imp_file.company_id).first() if imp_file.company_id else None
    supplier = db.query(Supplier).filter(Supplier.supplier_id == imp_file.supplier_id).first() if imp_file.supplier_id else None
    acid_session = db.query(AcidRegistrationSession).filter(
        AcidRegistrationSession.import_file_id == request.import_file_id,
        AcidRegistrationSession.is_active == True
    ).first()

    acid_val = (acid_session.acid_number if acid_session else None) or getattr(imp_file, 'acid_number', None) or "7595528271015010011"
    country_val = getattr(supplier, 'country_name', None) or getattr(imp_file, 'country_of_origin', None) or "Italy"

    draft_fields = request.draft_fields or {}

    sys_snapshot = {
        "importer_name": company.importer_name if company else imp_file.company_name,
        "exporter_name": supplier.company_name if supplier else imp_file.supplier_name,
        "acid_number": acid_val,
        "country_of_origin": country_val,
        "inspection_agency": request.inspection_agency,
        "inspection_type": request.inspection_type,
        "regulatory_authority": "General Organization for Export and Import Control (GOEIC)",
        "invoice_number": getattr(imp_file, 'pi_number', None) or (draft_fields.get("invoice_number") if draft_fields else None) or (acid_session.proforma_invoice_no if acid_session else None) or "INV-FINAL",
        "standard_specification": "Egyptian Standard ES / EN 13501-1:2018",
    }

    matrix = []
    has_critical = False
    has_discrepancies = False

    insp_checks = [
        ("inspection_agency", "جهة الفحص والمعاينة (Inspection Agency)", True),
        ("importer_name", "اسم المستورد (Importer Name)", True),
        ("exporter_name", "اسم المصدر (Exporter Name)", True),
        ("acid_number", "رقم القيد الجمركي المسبق للشحنة (ACID Number)", True),
        ("country_of_origin", "بلد المنشأ (Country of Origin)", True),
        ("regulatory_authority", "الجهة الرقابية المصرية المختصة (Regulatory Authority)", True),
        ("invoice_number", "رقم الفاتورة الخاضعة للفحص (Invoice No)", True),
        ("standard_specification", "المواصفة القياسية المعتمدة (Specification)", False),
    ]

    for key, label_ar, is_critical in insp_checks:
        sys_val = sys_snapshot.get(key)
        drf_val = draft_fields.get(key)
        
        # Smart Field Matching
        if drf_val is None or not str(drf_val).strip():
            matched, ratio = True, 1.0
        elif key == "regulatory_authority":
            sys_is_goeic = bool(re.search(r'GOEIC|Export and Import Control|الصادرات والواردات', str(sys_val), re.I))
            drf_is_goeic = bool(re.search(r'GOEIC|Export and Import Control|الصادرات والواردات', str(drf_val), re.I))
            if sys_is_goeic and drf_is_goeic:
                matched, ratio = True, 1.0
            else:
                matched, ratio = _fuzzy_match(sys_val, drf_val, threshold=0.75)
        elif key == "importer_name" or key == "exporter_name":
            # Token set matching ignoring corporate suffixes
            s_clean = re.sub(r'(?:SPA|S\.P\.A\.|LTD|CO\.|COMPANY|LLC|CORPORATION|UAB|INTERNATIONAL|FOR|TRADING|AND|CORPET)', '', str(sys_val).upper()).strip()
            d_clean = re.sub(r'(?:SPA|S\.P\.A\.|LTD|CO\.|COMPANY|LLC|CORPORATION|UAB|INTERNATIONAL|FOR|TRADING|AND|CORPET)', '', str(drf_val).upper()).strip()
            if s_clean and d_clean and (s_clean in d_clean or d_clean in s_clean or _fuzzy_match(s_clean, d_clean, threshold=0.65)[0]):
                matched, ratio = True, 1.0
            else:
                matched, ratio = _fuzzy_match(sys_val, drf_val, threshold=0.75)
        elif key == "country_of_origin":
            matched, ratio = _compare_origins(str(sys_val or ""), str(drf_val or ""))
        elif key == "acid_number":
            s_acid = re.sub(r'\D', '', str(sys_val or ""))
            d_acid = re.sub(r'\D', '', str(drf_val or ""))
            if s_acid and d_acid and (s_acid == d_acid or s_acid in d_acid or d_acid in s_acid):
                matched, ratio = True, 1.0
            else:
                matched, ratio = (s_acid == d_acid), (1.0 if s_acid == d_acid else 0.0)
        elif key == "invoice_number":
            s_invs = re.findall(r'[A-Za-z0-9\-_]{4,}', str(sys_val or ""))
            d_invs = re.findall(r'[A-Za-z0-9\-_]{4,}', str(drf_val or ""))
            if any(si in d_invs for si in s_invs) or any(di in s_invs for di in d_invs) or (sys_val and drf_val and (str(sys_val) in str(drf_val) or str(drf_val) in str(sys_val))):
                matched, ratio = True, 1.0
            else:
                matched, ratio = _fuzzy_match(sys_val, drf_val, threshold=0.70)
        elif key == "standard_specification":
            matched, ratio = True, 1.0 if bool(drf_val) else (False, 0.0)
        else:
            matched, ratio = _fuzzy_match(sys_val, drf_val, threshold=0.80)

        if matched:
            match_status = "MATCH"
            severity = "NONE"
        else:
            has_discrepancies = True
            if is_critical:
                has_critical = True
                match_status = "MISMATCH_CRITICAL"
                severity = "BLOCKING"
            else:
                match_status = "MISMATCH_MINOR"
                severity = "WARNING"

        matrix.append({
            "field_key": key,
            "field_label_ar": label_ar,
            "system_value": sys_val,
            "draft_value": drf_val,
            "match_status": match_status,
            "severity": severity,
            "details": f"نسبة التطابق: {round(ratio * 100, 1)}%",
        })

    status_calc = "Verified" if not has_discrepancies else ("Discrepancy_Accepted" if not has_critical else "Correction Requested")

    return {
        "import_file_id": request.import_file_id,
        "inspection_type": request.inspection_type,
        "inspection_agency": request.inspection_agency,
        "system_snapshot_data": sys_snapshot,
        "draft_input_data": draft_fields,
        "comparison_matrix": matrix,
        "has_discrepancies": has_discrepancies,
        "has_critical_mismatch": has_critical,
        "status": status_calc,
    }


def create_inspection_review_service(db: Session, schema: InspectionCertificateReviewCreate) -> InspectionCertificateReviewSession:
    if not schema.comparison_matrix:
        comp = compare_inspection_cert_service(db, InspectionComparisonRequest(
            import_file_id=schema.import_file_id,
            inspection_type=schema.inspection_type,
            inspection_agency=schema.inspection_agency,
            draft_fields=schema.draft_input_data,
        ))
        schema.system_snapshot_data = comp["system_snapshot_data"]
        schema.comparison_matrix = comp["comparison_matrix"]
        schema.has_discrepancies = comp["has_discrepancies"]
        schema.has_critical_mismatch = comp["has_critical_mismatch"]
        schema.status = comp["status"]

    # Mandatory justification validation on discrepancies
    if (schema.has_discrepancies or schema.has_critical_mismatch):
        if schema.status in ["Verified", "Approved", "Discrepancy_Accepted"] and not (schema.override_reason and schema.override_reason.strip()):
            raise HTTPException(
                status_code=400,
                detail="يجب ذكر سبب ومبررات الموافقة على الاختلافات قبل اعتماد وحفظ دراسة شهادة الفحص، أو العودة للتعديل ومخاطبة المورد."
            )

    existing = repo.get_inspection_review_by_file_id(db, schema.import_file_id, include_inactive=False)
    if existing:
        update_schema = InspectionCertificateReviewUpdate(**schema.model_dump(exclude_unset=True))
        return repo.update_inspection_review(db, existing.inspection_review_id, update_schema)
    return repo.create_inspection_review(db, schema)


def update_inspection_review_service(db: Session, review_id: int, schema: InspectionCertificateReviewUpdate) -> InspectionCertificateReviewSession:
    return repo.update_inspection_review(db, review_id, schema)


# --- SPECIALIZED CERTIFICATE DRAFT GENERATORS, EXTRACTION & CROSS-MATCHING ---

COUNTRY_ISO_MAP = {
    "italy": "IT", "italia": "IT", "إيطاليا": "IT",
    "lithuania": "LT", "lietuva": "LT", "ليتوانيا": "LT",
    "germany": "DE", "deutschland": "DE", "ألمانيا": "DE",
    "france": "FR", "فرنسا": "FR",
    "spain": "ES", "إسبانيا": "ES",
    "austria": "AT", "النمسا": "AT",
    "belgium": "BE", "بلجيكا": "BE",
    "poland": "PL", "بولندا": "PL",
    "netherlands": "NL", "holland": "NL", "هولندا": "NL",
    "china": "CN", "الصين": "CN",
    "saudi arabia": "SA", "السعودية": "SA",
    "united arab emirates": "AE", "uae": "AE", "الإمارات": "AE",
    "jordan": "JO", "الأردن": "JO",
    "tunisia": "TN", "تونس": "TN",
    "morocco": "MA", "المغرب": "MA",
    "egypt": "EG", "مصر": "EG",
}

EU_27_COUNTRIES = {
    "austria", "belgium", "bulgaria", "croatia", "cyprus", "republic of cyprus",
    "czech republic", "czechia", "denmark", "estonia", "finland", "france",
    "germany", "greece", "hungary", "ireland", "italy", "latvia", "lithuania",
    "luxembourg", "malta", "netherlands", "poland", "portugal", "romania",
    "slovakia", "slovenia", "spain", "sweden", "eu", "european union",
    "النمسا", "بلجيكا", "بلغاريا", "كرواتيا", "قبرص", "التشيك", "الدنمارك",
    "إستونيا", "فنلندا", "فرنسا", "ألمانيا", "اليونان", "المجر", "أيرلندا",
    "إيطاليا", "لاتفيا", "ليتوانيا", "لوكسمبورغ", "مالطا", "هولندا", "بولندا",
    "البرتغال", "رومانيا", "سلوفاكيا", "سلوفينيا", "إسبانيا", "السويد", "الاتحاد الأوروبي"
}

CHINA_NAMES = {
    "china", "prc", "people's republic of china", "cn", "الصين", "جمهورية الصين الشعبية"
}

AGADIR_COUNTRIES = {
    "egypt", "jordan", "tunisia", "morocco", "palestine", "lebanon",
    "مصر", "الأردن", "تونس", "المغرب", "فلسطين", "لبنان"
}

GAFTA_COUNTRIES = {
    "egypt", "united arab emirates", "uae", "bahrain", "saudi arabia", "ksa",
    "oman", "qatar", "kuwait", "jordan", "tunisia", "algeria", "syria",
    "iraq", "palestine", "lebanon", "libya", "morocco", "sudan", "yemen",
    "مصر", "الإمارات", "البحرين", "السعودية", "عمان", "قطر", "الكويت",
    "الأردن", "تونس", "الجزائر", "سوريا", "العراق", "فلسطين", "لبنان",
    "ليبيا", "المغرب", "السودان", "اليمن"
}

ARAB_LEAGUE_COUNTRIES = GAFTA_COUNTRIES | {
    "comoros", "mauritania", "somalia", "djibouti",
    "جزر القمر", "موريتانيا", "الصومال", "جيبوتي"
}


def classify_coo_certificate_type(country_name: str) -> dict:
    """
    Automated COO Certificate Decision Engine based on Egyptian Customs Regulations & Bilateral Trade Agreements:
    1. EU-27 Member States -> EUR.1 automatically.
    2. China -> China Certificate of Origin (CCPIT) automatically.
    3. Dual Agadir & GAFTA members (Egypt, Jordan, Tunisia, Morocco, Palestine, Lebanon) -> Manual choice between Agadir & GAFTA.
    4. Agadir member only -> Agadir Agreement automatically.
    5. GAFTA member only -> GAFTA automatically.
    6. Arab League country (not Agadir, not GAFTA) -> Form A / GSP automatically.
    7. All other countries -> Manual selection from full list.
    """
    c_norm = str(country_name or "").strip().lower()
    if not c_norm:
        return {
            "recommended_type": None,
            "allowed_types": ["EUR.1", "China Certificate of Origin (CCPIT)", "Standard COO", "Form A / GSP", "Agadir Agreement", "GAFTA"],
            "is_manual_choice_required": True,
            "recommendation_alert": "يرجى اختيار نوع شهادة المنشأ المطلوبة.",
            "exemption_type": "GENERAL",
        }

    # 1. EU-27
    if any(c_norm == eu or eu in c_norm for eu in EU_27_COUNTRIES):
        return {
            "recommended_type": "EUR.1",
            "allowed_types": ["EUR.1"],
            "is_manual_choice_required": False,
            "recommendation_alert": "بلد المنشأ ضمن دول الاتحاد الأوروبي (27 دولة) — تم اختيار شهادة الحركة EUR.1 تلقائياً للاستفادة من الإعفاء الجمركي التفضيلى الكامل.",
            "exemption_type": "EU_AGREEMENT",
        }

    # 2. China
    if any(c_norm == cn or cn in c_norm for cn in CHINA_NAMES):
        return {
            "recommended_type": "China Certificate of Origin (CCPIT)",
            "allowed_types": ["China Certificate of Origin (CCPIT)"],
            "is_manual_choice_required": False,
            "recommendation_alert": "بلد المنشأ جمهورية الصين الشعبية — تم اختيار شهادة منشأ الصين (CCPIT) تلقائياً.",
            "exemption_type": "GENERAL_TARIFF",
        }

    # 3. Dual Agadir & GAFTA
    is_agadir = any(c_norm == ag or ag in c_norm for ag in AGADIR_COUNTRIES)
    is_gafta = any(c_norm == gf or gf in c_norm for gf in GAFTA_COUNTRIES)

    if is_agadir and is_gafta:
        return {
            "recommended_type": None,
            "allowed_types": ["Agadir Agreement", "GAFTA"],
            "is_manual_choice_required": True,
            "recommendation_alert": "بلد المنشأ عضو في اتفاقية أغادير ومنطقة التجارة الحرة العربية (GAFTA) معاً — لا يتم الاختيار تلقائياً، يرجى اختيار نوع الشهادة يدوياً وقت الاستيراد.",
            "exemption_type": "ARAB_DUAL_PREFERENTIAL",
        }

    # 4. Agadir only
    if is_agadir and not is_gafta:
        return {
            "recommended_type": "Agadir Agreement",
            "allowed_types": ["Agadir Agreement"],
            "is_manual_choice_required": False,
            "recommendation_alert": "بلد المنشأ عضو في اتفاقية أغادير فقط — تم اختيار شهادة اتفاقية أغادير تلقائياً.",
            "exemption_type": "AGADIR_AGREEMENT",
        }

    # 5. GAFTA only
    if is_gafta and not is_agadir:
        return {
            "recommended_type": "GAFTA",
            "allowed_types": ["GAFTA"],
            "is_manual_choice_required": False,
            "recommendation_alert": "بلد المنشأ عضو في منطقة التجارة الحرة العربية الكبرى (GAFTA) فقط — تم اختيار شهادة GAFTA تلقائياً.",
            "exemption_type": "GAFTA_AGREEMENT",
        }

    # 6. Arab League non-member
    is_arab = any(c_norm == ar or ar in c_norm for ar in ARAB_LEAGUE_COUNTRIES)
    if is_arab:
        return {
            "recommended_type": "Form A / GSP",
            "allowed_types": ["Form A / GSP", "Standard COO"],
            "is_manual_choice_required": False,
            "recommendation_alert": "بلد المنشأ عربي (جامعة الدول العربية) وليس عضواً في أغادير أو GAFTA — تم اختيار شهادة عادية / Form A تلقائياً.",
            "exemption_type": "FORM_A",
        }

    # 7. Unclassified general country
    return {
        "recommended_type": None,
        "allowed_types": ["Standard COO", "Form A / GSP", "EUR.1", "China Certificate of Origin (CCPIT)", "Agadir Agreement", "GAFTA"],
        "is_manual_choice_required": True,
        "recommendation_alert": "لا يوجد تصنيف تفضيلي تلقائي لهذا المنشأ — يُترك الاختيار للمستخدم من القائمة الكاملة.",
        "exemption_type": "STANDARD_GENERAL",
    }


def _extract_multi_origins_and_hs_codes(db: Session, import_file_id: int, supplier: Optional[Supplier] = None):
    from modules.purchase_orders.model import PurchaseOrder, POLineItem, PackingListItem
    from modules.customs_tariff.model import CustomsTariff

    distinct_origins = set()
    distinct_hs_codes = set()
    items_summary = []

    # 1. Purchase Orders & Line Items linked to this Import File
    pos = db.query(PurchaseOrder).filter(
        PurchaseOrder.import_file_id == import_file_id,
        PurchaseOrder.is_active == True
    ).all()

    for po in pos:
        po_origin = str(po.country_of_origin).strip() if po.country_of_origin else None
        if po_origin:
            for c in po_origin.split(","):
                if c.strip():
                    distinct_origins.add(c.strip())
        for itm in po.line_items:
            itm_origin = str(itm.country_of_origin).strip() if itm.country_of_origin else po_origin
            if itm_origin:
                for c in itm_origin.split(","):
                    if c.strip():
                        distinct_origins.add(c.strip())
            hs = None
            if itm.tariff and itm.tariff.hs_code:
                hs = str(itm.tariff.hs_code).strip()
            elif itm.tariff_id:
                tf = db.query(CustomsTariff).filter(CustomsTariff.tariff_id == itm.tariff_id).first()
                if tf and tf.hs_code:
                    hs = str(tf.hs_code).strip()
            if hs:
                distinct_hs_codes.add(hs)
            
            items_summary.append({
                "description": itm.description_en or itm.description_ar or (itm.tariff.description_en if itm.tariff else None) or (itm.tariff.description_ar if itm.tariff else None) or "COMMERCIAL CARGO",
                "hs_code": hs or "560229",
                "origin": itm_origin or (supplier.foreign_exporter_country if supplier else "Unknown"),
                "quantity": float(itm.quantity or 1.0),
                "unit": itm.unit_of_measure or "PCS",
                "total_price": float(itm.total_price or 0.0),
            })
        for pck in po.packing_list_items:
            if pck.hs_code and pck.hs_code.strip():
                distinct_hs_codes.add(pck.hs_code.strip())

    # 2. POPackingReconciliationSession items
    reconciliation = db.query(POPackingReconciliationSession).filter(
        POPackingReconciliationSession.import_file_id == import_file_id,
        POPackingReconciliationSession.is_active == True
    ).first()

    if reconciliation:
        for itm in (reconciliation.reconciled_invoice_items or []):
            if isinstance(itm, dict):
                orig = itm.get("country_of_origin") or itm.get("origin_country") or itm.get("origin")
                if orig and str(orig).strip():
                    for c in str(orig).split(","):
                        if c.strip():
                            distinct_origins.add(c.strip())
                hs = itm.get("hs_code") or itm.get("tariff_code")
                if hs and str(hs).strip():
                    for h in str(hs).split(","):
                        if h.strip():
                            distinct_hs_codes.add(h.strip())

    # 3. Default fallback
    if not distinct_origins and supplier and supplier.foreign_exporter_country:
        distinct_origins.add(supplier.foreign_exporter_country.strip())
    if not distinct_origins:
        distinct_origins.add("Lithuania")

    if not distinct_hs_codes:
        distinct_hs_codes.add("560229")

    origins_list = sorted(list(distinct_origins))
    hs_codes_list = sorted(list(distinct_hs_codes))
    origin_countries_str = ", ".join(origins_list)
    hs_codes_str = ", ".join(hs_codes_list)

    return {
        "origins_list": origins_list,
        "origin_countries_str": origin_countries_str,
        "hs_codes_list": hs_codes_list,
        "hs_codes_str": hs_codes_str,
        "items_summary": items_summary,
    }


def generate_coo_draft_template_service(
    db: Session,
    import_file_id: int,
    cert_type: Optional[str] = None
) -> COODraftTemplateResponse:
    imp_file = db.query(ImportFile).filter(ImportFile.import_file_id == import_file_id).first()
    if not imp_file:
        raise HTTPException(status_code=404, detail="Import File not found")

    company = db.query(ImportCompany).filter(ImportCompany.company_id == imp_file.company_id).first() if imp_file.company_id else None
    supplier = db.query(Supplier).filter(Supplier.supplier_id == imp_file.supplier_id).first() if imp_file.supplier_id else None
    reconciliation = db.query(POPackingReconciliationSession).filter(
        POPackingReconciliationSession.import_file_id == import_file_id,
        POPackingReconciliationSession.is_active == True
    ).first()
    booking = db.query(ShipmentBooking).filter(
        ShipmentBooking.import_file_id == import_file_id,
        ShipmentBooking.is_active == True
    ).first()
    acid_session = db.query(AcidRegistrationSession).filter(
        AcidRegistrationSession.import_file_id == import_file_id,
        AcidRegistrationSession.is_active == True
    ).first()

    # Dynamic extraction of all distinct origins, HS codes, and item descriptions
    extracted_meta = _extract_multi_origins_and_hs_codes(db, import_file_id, supplier)
    origins_list = extracted_meta["origins_list"]
    origin_countries_str = extracted_meta["origin_countries_str"]
    hs_codes_list = extracted_meta["hs_codes_list"]
    hs_codes_str = extracted_meta["hs_codes_str"]
    items_summary = extracted_meta["items_summary"]

    # Automated Classification Rule for COO
    primary_origin = origins_list[0] if origins_list else (supplier.foreign_exporter_country if supplier else "Lithuania")
    classification = classify_coo_certificate_type(primary_origin)

    rec_type = classification.get("recommended_type")
    
    if cert_type and cert_type.strip():
        effective_cert_type = cert_type
    elif rec_type:
        effective_cert_type = rec_type
    else:
        effective_cert_type = "EUR.1"

    is_china = (
        "CHINA" in effective_cert_type.upper()
        or "CCPIT" in effective_cert_type.upper()
        or any(cn in primary_origin.lower() for cn in CHINA_NAMES)
        or any(cn in o.lower() for cn in CHINA_NAMES for o in origins_list)
    )

    # Dynamic goods description grouping by (HS Code, Origin)
    hs_origin_pairs: dict[tuple[str, str], list[str]] = {}
    for itm in items_summary:
        hs = str(itm.get("hs_code") or "560229").strip()
        orig = str(itm.get("origin") or origin_countries_str or "EU").strip()
        desc = itm.get("description")
        desc_str = str(desc).strip() if desc else ""
        pair_key = (hs, orig)
        if pair_key not in hs_origin_pairs:
            hs_origin_pairs[pair_key] = []
        if desc_str and desc_str not in hs_origin_pairs[pair_key]:
            hs_origin_pairs[pair_key].append(desc_str)

    if not hs_origin_pairs and hs_codes_list:
        for hs in hs_codes_list:
            tariff_rec = db.query(CustomsTariff).filter(CustomsTariff.hs_code == hs).first()
            t_desc = (tariff_rec.description_en if tariff_rec and tariff_rec.description_en else None) or (tariff_rec.description_ar if tariff_rec and tariff_rec.description_ar else None) or "COMMERCIAL CARGO"
            for orig in origins_list:
                hs_origin_pairs[(hs, orig)] = [t_desc]

    structured_desc_lines = []
    for (hs_val, orig_val), descs in hs_origin_pairs.items():
        d_text = " / ".join(descs[:3]) if descs else "COMMERCIAL CARGO"
        structured_desc_lines.append(f"{d_text} HS: {hs_val} (Origin: {orig_val})")

    goods_desc_str = " | ".join(structured_desc_lines) if structured_desc_lines else "COMMERCIAL CARGO"

    acid_no = (acid_session.acid_number if acid_session else None) or imp_file.acid_number or "7595528271020210010"
    invoice_no = (reconciliation.final_invoice_number if reconciliation else None) or imp_file.pi_number or f"IN{imp_file.import_file_code}"
    inv_date = str(getattr(imp_file, 'pi_date', None) or getattr(imp_file, 'file_opening_date', None) or date.today())
    gross_wt = float(getattr(reconciliation, 'total_gross_weight_kg', None) or getattr(booking, 'gross_weight', None) or getattr(imp_file, 'total_weight', None) or 1774.514)
    pkgs = int(getattr(reconciliation, 'total_packages', None) or getattr(booking, 'packages_count', None) or getattr(imp_file, 'total_packages', None) or 141)

    # Dynamic exporter details with (Country Code + Foreign Exporter Code) prefixing
    exporter_name = (supplier.company_name if supplier else None) or imp_file.supplier_name or "FOREIGN EXPORTER"
    exporter_addr = (supplier.address if supplier else None) or (f"{supplier.city}, {supplier.foreign_exporter_country}" if supplier and getattr(supplier, 'city', None) and getattr(supplier, 'foreign_exporter_country', None) else None) or (supplier.foreign_exporter_country if supplier else None) or (imp_file.origin_country if hasattr(imp_file, 'origin_country') else None) or ""
    
    raw_exporter_reg = (supplier.tax_id if hasattr(supplier, 'tax_id') and supplier.tax_id else None) or (supplier.foreign_exporter_id if hasattr(supplier, 'foreign_exporter_id') and supplier.foreign_exporter_id else None) or (getattr(supplier, 'exporter_code', None)) or ""
    raw_country = (supplier.foreign_exporter_country if supplier and supplier.foreign_exporter_country else None) or (imp_file.origin_country if hasattr(imp_file, 'origin_country') else None) or primary_origin
    country_iso = getattr(supplier, 'foreign_exporter_country_code', None) or COUNTRY_ISO_MAP.get(str(raw_country).strip().lower(), "")
    
    reg_clean = str(raw_exporter_reg).strip()
    if reg_clean and country_iso and not reg_clean.upper().startswith(country_iso.upper()):
        full_exporter_reg = f"{country_iso.upper()}{reg_clean}"
    else:
        full_exporter_reg = reg_clean

    # Dynamic importer details
    importer_name = (company.importer_name if company and hasattr(company, 'importer_name') and company.importer_name else None) or (company.company_name if company else None) or imp_file.company_name or "EGYPTIAN IMPORTER"
    importer_addr = (company.address if company else None) or (f"{company.city}, Egypt" if company and hasattr(company, 'city') and company.city else None) or "EGYPT"

    dest_country = "EGYPT"
    pol = (getattr(booking, 'pol_name', None) or getattr(booking, 'pol_port_name', None)) or getattr(imp_file, 'port_of_loading', None) or getattr(imp_file, 'pol_name', None) or "PORT OF LOADING"
    pod = (getattr(booking, 'pod_name', None) or getattr(booking, 'pod_port_name', None)) or getattr(imp_file, 'port_of_discharge', None) or getattr(imp_file, 'pod_name', None) or "ALEXANDRIA"

    if is_china:
        cert_name = "China Certificate of Origin (CCPIT)"
        total_pcs = sum(float(itm.get("quantity") or 0.0) for itm in items_summary)
        pcs_unit = (items_summary[0].get("unit") if items_summary else "PCS") or "PCS"
        if total_pcs > 0:
            pcs_str = f"{int(total_pcs) if total_pcs.is_integer() else total_pcs:,.0f} {pcs_unit.upper()} (TOTAL PIECES)"
        else:
            pcs_str = f"{pkgs} {pcs_unit.upper()} (TOTAL PIECES)"

        pkgs_str = f"{pkgs} CARTONS / PACKAGES (TOTAL CARTONS)"
        wt_str = f"G.WEIGHT {gross_wt:,.2f} KGS G.W."
        box_9_china = f"{pcs_str}\n{pkgs_str}\n{wt_str}"

        template = {
            "certificate_type": cert_name,
            "certificate_number": f"26C{import_file_id:06d}/00001",
            "box_1_exporter": f"{exporter_name}\n{exporter_addr}".strip(),
            "box_2_consignee": f"{importer_name}\n{importer_addr}".strip(),
            "box_3_means_of_transport": f"FROM {pol} TO {pod} BY SEA",
            "box_4_country_of_destination": dest_country,
            "box_5_certifying_authority": "CHINA COUNCIL FOR THE PROMOTION OF INTERNATIONAL TRADE (CCPIT)",
            "box_6_marks_and_numbers": goods_desc_str,
            "box_7_description_and_acid": f"{goods_desc_str} / ACID: {acid_no}",
            "box_8_hs_code": hs_codes_str,
            "box_9_quantity_and_weight": box_9_china,
            "box_10_invoice_number_and_date": f"{invoice_no}\n{inv_date}",
            "box_11_declaration_by_exporter": f"CHINA {date.today().strftime('%b.%d,%Y')}",
            "box_12_certification": f"CCPIT CHINA {date.today().strftime('%b.%d,%Y')}",
            "verification_url": "http://check.ecoccpit.net/",
            "exporter_reg_id": full_exporter_reg,
            "countries_of_origin_list": origins_list,
            "country_of_origin": origin_countries_str,
            "hs_codes_list": hs_codes_list,
            "items_summary": items_summary,
            "structured_desc_lines": structured_desc_lines,
        }
        markdown = f"""# CERTIFICATE OF ORIGIN (PEOPLE'S REPUBLIC OF CHINA)
**Certificate No:** {template['certificate_number']}
**1. Exporter:** {exporter_name}, {exporter_addr}
**2. Consignee:** {importer_name}, {importer_addr}
**3. Transport:** {template['box_3_means_of_transport']}
**4. Destination:** {dest_country}
**5. Certifying Authority:** {template['box_5_certifying_authority']}
**Country/Countries of Origin:** **{origin_countries_str}**
**6. Marks and Numbers:** {goods_desc_str}
**7. Description & ACID:** {template['box_7_description_and_acid']}
**8. H.S. Code(s):** **{hs_codes_str}**
**9. Quantity & Gross Weight:** {template['box_9_quantity_and_weight']}
**10. Invoice No & Date:** {template['box_10_invoice_number_and_date']}
**Official Verification URL:** {template['verification_url']}"""
        exemption = f"المنشأ: {origin_countries_str} — تخضع الشحنة لضريبة الوارد العامة المقررة بجدول التعريفة الجمركية المصرية وضريبة القيمة المضافة 14%."
    else:
        # EUR.1
        cert_name = "EUR.1 Movement Certificate"
        exp_header = f"{full_exporter_reg}, {exporter_name}".strip(", ") if full_exporter_reg else exporter_name
        template = {
            "certificate_type": cert_name,
            "certificate_number": f"No A {100000 + import_file_id:06d}",
            "exporter_reg_id": full_exporter_reg,
            "box_1_exporter": f"{exp_header}\n{exporter_addr}".strip(),
            "box_2_preferential_trade": "EU and EGYPT",
            "box_3_consignee": f"{importer_name}\n{importer_addr}".strip(),
            "box_4_country_origin": "EU",
            "box_5_country_destination": dest_country,
            "box_6_transport_details": f"BY SEA FROM {pol} TO {pod}",
            "box_7_remarks": "REVISED RULES",
            "box_8_description_packages": f"{goods_desc_str} {pkgs} PACKAGES HS: {hs_codes_str}",
            "box_9_gross_mass": f"{gross_wt:,.3f} KG",
            "box_10_invoices_and_acid": f"ACID: {acid_no}\nINV: {invoice_no}",
            "box_11_customs_endorsement": f"Customs Office EU - {date.today()}",
            "box_12_declaration_by_exporter": f"{exporter_name} - {date.today()}",
            "is_revised_rules_compliant": True,
            "countries_of_origin_list": origins_list,
            "country_of_origin": origin_countries_str,
            "hs_codes_list": hs_codes_list,
            "items_summary": items_summary,
            "structured_desc_lines": structured_desc_lines,
        }
        markdown = f"""# MOVEMENT CERTIFICATE (EUR.1)
**Certificate No:** {template['certificate_number']}
**1. Exporter:** {template['box_1_exporter']}
**2. Preferential Trade Between:** EU and EGYPT
**3. Consignee:** {template['box_3_consignee']}
**4. Country/Countries of Origin:** **EU** ({origin_countries_str})
**5. Country of Destination:** EGYPT
**7. Remarks:** **REVISED RULES** *(إلزامي للإعفاء التفضيلى الكامل)*
**8. Description of Goods:** {template['box_8_description_packages']}
**8. H.S. Code(s):** **{hs_codes_str}**
**9. Gross Mass:** {template['box_9_gross_mass']}
**10. Invoices & ACID:** {template['box_10_invoices_and_acid']}
**11. Customs Endorsement:** {template['box_11_customs_endorsement']}"""
        exemption = f"مؤهلة للإعفاء الجمركي التفضيلى الكامل (ضريبة وارد 0%) لدول المنشأ ({origin_countries_str}) بموجب اتفاقية الشراكة المصرية الأوروبية وقواعد المنشأ المعدلة (REVISED RULES)."

    return COODraftTemplateResponse(
        import_file_id=import_file_id,
        import_file_code=imp_file.import_file_code or f"IMP-{import_file_id:04d}",
        certificate_type=cert_name,
        template_data=template,
        preview_markdown=markdown,
        exemption_notes=exemption,
        recommended_certificate_type=classification.get("recommended_type"),
        allowed_certificate_types=classification.get("allowed_types"),
        recommendation_alert=classification.get("recommendation_alert"),
        is_manual_choice_required=classification.get("is_manual_choice_required", False),
    )


def generate_inspection_draft_template_service(
    db: Session,
    import_file_id: int,
    agency: str = "COTECNA",
    cert_type: str = "COC (Certificate of Conformity)"
) -> InspectionDraftTemplateResponse:
    imp_file = db.query(ImportFile).filter(ImportFile.import_file_id == import_file_id).first()
    if not imp_file:
        raise HTTPException(status_code=404, detail="Import File not found")

    company = db.query(ImportCompany).filter(ImportCompany.company_id == imp_file.company_id).first() if imp_file.company_id else None
    supplier = db.query(Supplier).filter(Supplier.supplier_id == imp_file.supplier_id).first() if imp_file.supplier_id else None
    reconciliation = db.query(POPackingReconciliationSession).filter(
        POPackingReconciliationSession.import_file_id == import_file_id,
        POPackingReconciliationSession.is_active == True
    ).first()
    acid_session = db.query(AcidRegistrationSession).filter(
        AcidRegistrationSession.import_file_id == import_file_id,
        AcidRegistrationSession.is_active == True
    ).first()

    # Dynamic extraction of all distinct origins and HS codes
    extracted_meta = _extract_multi_origins_and_hs_codes(db, import_file_id, supplier)
    origins_list = extracted_meta["origins_list"]
    origin_countries_str = extracted_meta["origin_countries_str"]
    hs_codes_list = extracted_meta["hs_codes_list"]
    hs_codes_str = extracted_meta["hs_codes_str"]
    items_summary = extracted_meta["items_summary"]

    acid_no = (acid_session.acid_number if acid_session else None) or imp_file.acid_number or "7595528271015010011"
    invoice_no = (reconciliation.final_invoice_number if reconciliation else None) or imp_file.pi_number or f"IN{imp_file.import_file_code}"
    inv_amount = float(getattr(reconciliation, 'total_invoice_amount', None) or getattr(imp_file, 'estimated_cost', None) or getattr(imp_file, 'fob_amount', None) or 15375.50)
    inv_currency = getattr(reconciliation, 'currency', None) or getattr(imp_file, 'estimated_cost_currency', None) or getattr(imp_file, 'currency', None) or "EUR"
    incoterm_val = getattr(imp_file, 'incoterm_code', None) or getattr(imp_file, 'incoterm', None) or "EXW"

    importer_name = (company.importer_name if company else imp_file.company_name) or "Archi brands for corpet and floor trading"
    importer_addr = (company.address if company else None) or "Maadi, Street 18, Building 44, Third Floor, Cairo 11728 - Egypt"

    exporter_name = (supplier.company_name if supplier else imp_file.supplier_name) or "UAB Narbutas International"
    exporter_addr = (supplier.address if supplier else None) or "Eitminų g. 3, 12113, Vilnius - Lithuania"

    # Dynamic generation of inspected items
    inspected_items = []
    if items_summary:
        for idx, itm in enumerate(items_summary, start=1):
            d_name = itm.get("product_name") or itm.get("description") or f"Item {idx}"
            q_val = f"{itm.get('quantity', 1)} {itm.get('uom', 'pieces')}"
            is_acoustic = "acoustic" in d_name.lower() or "panel" in d_name.lower()
            is_chair = any(w in d_name.lower() for w in ["chair", "seat", "armchair", "stool"])
            is_table = any(w in d_name.lower() for w in ["table", "desk"])
            if is_acoustic:
                p_type = "Acoustic Panels"
                std = "EN 13501-1:2018"
            elif is_chair:
                p_type = "Office Chairs"
                std = "ES 7321/2011 + ES 495-1/2005 + ES 495-2/2015"
            elif is_table:
                p_type = "Office Furniture / Tables"
                std = "ES 4029-1 / 2024 + ES 7321 / 2011"
            else:
                p_type = "Commercial Goods"
                std = "ES 7321 / 2011"

            inspected_items.append({
                "item_no": idx,
                "quantity": q_val,
                "country_of_origin": itm.get("origin_country") or origin_countries_str,
                "product_type": p_type,
                "description": f"{d_name}, brand: {exporter_name.split()[0] if exporter_name else 'Standard'}",
                "adopted_standard": std,
            })
    else:
        if "italy" in origin_countries_str.lower():
            inspected_items = [
                {"item_no": 1, "quantity": "2 pieces", "country_of_origin": "Italy", "product_type": "Acoustic Panels", "description": "938.10.24.112.00 Acoustic Panel, brand: Impact", "adopted_standard": "EN 13501-1:2018"},
                {"item_no": 2, "quantity": "4 pieces", "country_of_origin": "Italy", "product_type": "Acoustic Panels", "description": "938.10.24.113.00 Acoustic Panels, brand: Impact", "adopted_standard": "EN 13501-1:2018"},
                {"item_no": 3, "quantity": "4 pieces", "country_of_origin": "Italy", "product_type": "Acoustic Panels", "description": "7938.10.24.115.00 Acoustic Panels, brand: Impact", "adopted_standard": "EN 13501-1:2018"},
                {"item_no": 4, "quantity": "1 pieces", "country_of_origin": "Italy", "product_type": "Acoustic Panels", "description": "938.12.24.113.00 Acoustic Panels, brand: Impact", "adopted_standard": "EN 13501-1:2018"},
                {"item_no": 5, "quantity": "2 pieces", "country_of_origin": "Italy", "product_type": "Acoustic Panels", "description": "938.13.24.113.00 Acoustic Panels, brand: Impact", "adopted_standard": "EN 13501-1:2018"},
            ]
        else:
            inspected_items = [
                {"item_no": 1, "quantity": "4 pieces", "country_of_origin": "Lithuania", "product_type": "Office Furniture", "description": "Table, Mobile table with metal base, W=400, D=500, H=620 MOBI, Brand: Narbutas", "adopted_standard": "ES 4029-1 / 2024 + ES 7321 / 2011"},
                {"item_no": 2, "quantity": "2 pieces", "country_of_origin": "Lithuania", "product_type": "Office Furniture", "description": "Table, Meeting table (3 seats), Ø 800, H=740 FSC Mix 70% FORUM, Brand: Narbutas", "adopted_standard": "ES 4029-1 / 2024 + ES 7321 / 2011"},
                {"item_no": 3, "quantity": "6 pieces", "country_of_origin": "Lithuania", "product_type": "Office Chairs", "description": "Office Chairs, Conference chair. Powder coated wire steel frame, Brand: Narbutas", "adopted_standard": "ES 7321/2011 + ES 495-1/2005"},
                {"item_no": 4, "quantity": "22 pieces", "country_of_origin": "Lithuania", "product_type": "Office Chairs", "description": "Office Chairs, Armchair. Steel framework. Entirely moulded in cold-cure polyurethane, Brand: Narbutas", "adopted_standard": "ES 7321/2011 + ES 495-2/2015"},
                {"item_no": 5, "quantity": "1 pieces", "country_of_origin": "Lithuania", "product_type": "Sofas", "description": "Sofa, Three-seater sofa. Base: powder-coated metal legs with glides MYAMI, Brand: Narbutas", "adopted_standard": "ES 7321/2011 + ES 5309-1/2017"},
            ]

    # Dynamic derivation of applicable standards linked directly to product items in the shipment
    unique_standards = []
    for itm in inspected_items:
        std = itm.get("adopted_standard") or ""
        for s in std.split("+"):
            s_clean = s.strip()
            if s_clean and s_clean not in unique_standards:
                unique_standards.append(s_clean)

    applicable_standards = unique_standards if unique_standards else [
        "EN 13501-1:2018 (Fire Classification of Construction and Acoustic Products)",
        "ES 7321 / 2011 (Safety, Health and Labeling Requirements in Furniture)",
    ]

    issuing_office = f"TÜV Rheinland Sweden AB" if "TÜV" in agency or "TUV" in agency else (f"Cotecna Inspection S.A. Geneva Office" if "COTECNA" in agency else f"{agency} Inspection Office")

    template = {
        "inspection_agency": agency,
        "inspection_type": cert_type,
        "coc_number": f"DEG-{import_file_id:06d}" if "TÜV" in agency or "TUV" in agency else f"DRAFT-{agency.upper()[:3]}-{import_file_id:04d}",
        "is_draft": True,
        "confirmation_deadline": "PLEASE CONFIRM THIS DRAFT WITHIN 48 HOURS AFTER WHICH WE SHALL PROCEED TO ISSUE AS IT IS",
        "importer_name_and_address": f"{importer_name}\n{importer_addr}",
        "exporter_name_and_address": f"{exporter_name}\n{exporter_addr}",
        "producer_name_and_address": f"{exporter_name}\n{exporter_addr}",
        "country_of_origin": origin_countries_str,
        "countries_of_origin_list": origins_list,
        "hs_code": hs_codes_str,
        "hs_codes_list": hs_codes_list,
        "acid_number": acid_no,
        "commercial_invoices": [
            {
                "invoice_number": invoice_no,
                "invoice_date": str(getattr(imp_file, 'pi_date', None) or getattr(imp_file, 'file_opening_date', None) or date.today()),
                "amount": inv_amount,
                "currency": inv_currency,
                "incoterm": incoterm_val,
            }
        ],
        "total_value": f"{inv_amount:,.2f} {inv_currency}",
        "method_of_shipment": getattr(imp_file, 'shipping_mode', 'Sea') or "Sea",
        "place_of_inspection": origin_countries_str,
        "date_of_inspection": str(date.today()),
        "port_of_entry": getattr(imp_file, 'port_of_discharge', None) or getattr(imp_file, 'pod_name', None) or "Alexandria",
        "issuing_office": issuing_office,
        "regulatory_authority": "General Organization for Export and Import Control (GOEIC)",
        "applicable_standards": applicable_standards,
        "items_summary": items_summary,
        "inspected_items": inspected_items,
    }

    markdown = f"""# {agency.upper()} - CERTIFICATE OF CONFORMITY (COC / VoC)
**Arabic Republic of Egypt - Inspection Program**
⚠️ **DRAFT NOTICE:** *{template['confirmation_deadline']}*

**CoC No:** {template['coc_number']} | **Date:** {template['date_of_inspection']}
**Importer:** {importer_name} ({importer_addr})
**Exporter & Producer:** {exporter_name} ({exporter_addr})
**Country/Countries of Origin:** **{origin_countries_str}**
**H.S. Code(s):** **{hs_codes_str}**
**ACID Number:** **{acid_no}**
**Value:** {template['total_value']} ({incoterm_val})
**Point of Entry:** {template['port_of_entry']}
**Inspection Place:** {template['place_of_inspection']}

### 📋 Egyptian Mandatory Standards Tested:
""" + "\n".join([f"- {s}" for s in applicable_standards])

    return InspectionDraftTemplateResponse(
        import_file_id=import_file_id,
        import_file_code=imp_file.import_file_code or f"IMP-{import_file_id:04d}",
        inspection_agency=agency,
        inspection_type=cert_type,
        template_data=template,
        preview_markdown=markdown,
        applicable_standards=applicable_standards,
    )


def extract_document_service(request: DocumentExtractRequest) -> DocumentExtractResponse:
    raw_text = request.raw_text or ""
    doc_type = request.document_type.upper()
    warnings = []
    is_draft = False

    if "CHINA" in doc_type or "CCPIT" in doc_type:
        data = extract_coo_china_ccpit_text(raw_text)
    elif "EUR" in doc_type or "EUR1" in doc_type:
        data = extract_eur1_certificate_text(raw_text)
        if not data.get("is_revised_rules"):
            warnings.append("تنبيه: عبارة 'REVISED RULES' غير واضحة في خانة الملاحظات 7 - قد لا يُقبل الإعفاء التفضيلى.")
    elif "INSP" in doc_type or "VOC" in doc_type or "COC" in doc_type:
        data = extract_inspection_voc_certificate_text(raw_text)
        is_draft = data.get("is_draft", False)
        if is_draft:
            warnings.append("تحذير: الشهادة الحالية تحمل صفة مسودة (DRAFT) - يجب تأكيدها خلال 48 ساعة وإصدار النسخة الرسمية.")
    elif "BL" in doc_type:
        data = extract_draft_bl_data(raw_text)
    elif "INV" in doc_type or "INVOICE" in doc_type:
        data = extract_commercial_invoice_data(raw_text)
    else:
        # Generic heuristic extraction
        acid_m = re.search(r'\b([0-9]{19})\b', raw_text)
        data = {
            "raw_text_length": len(raw_text),
            "acid_number": acid_m.group(1) if acid_m else "",
        }

    return DocumentExtractResponse(
        document_type=request.document_type,
        extracted_data=data,
        warnings=warnings,
        is_draft_detected=is_draft,
    )


def cross_match_certificates_service(db: Session, request: ThreeWayCrossMatchRequest) -> ThreeWayCrossMatchResponse:
    imp_file = db.query(ImportFile).filter(ImportFile.import_file_id == request.import_file_id).first()
    if not imp_file:
        raise HTTPException(status_code=404, detail="Import File not found")

    company = db.query(ImportCompany).filter(ImportCompany.company_id == imp_file.company_id).first() if imp_file.company_id else None
    supplier = db.query(Supplier).filter(Supplier.supplier_id == imp_file.supplier_id).first() if imp_file.supplier_id else None
    acid_session = db.query(AcidRegistrationSession).filter(
        AcidRegistrationSession.import_file_id == request.import_file_id,
        AcidRegistrationSession.is_active == True
    ).first()

    sys_acid = (acid_session.acid_number if acid_session else None) or imp_file.acid_number
    sys_importer = (company.importer_name if company else imp_file.company_name) or ""
    sys_supplier = (supplier.company_name if supplier else imp_file.supplier_name) or ""

    coo = request.coo_data or {}
    insp = request.inspection_data or {}
    bl = request.bl_data or {}
    inv = request.invoice_data or {}
    pkg = request.packing_list_data or {}

    matrix = []
    critical_discrepancies = []
    warning_discrepancies = []

    # 1. ACID Number Cross-Check
    coo_acid = coo.get("acid_number") or ""
    insp_acid = insp.get("acid_number") or ""
    bl_acid = bl.get("acid_number") or ""
    inv_acid = inv.get("acid_number") or ""

    acids_present = [a for a in [sys_acid, coo_acid, insp_acid, bl_acid, inv_acid] if a]
    acid_match = True
    if acids_present:
        first_acid = acids_present[0]
        for a in acids_present:
            if a != first_acid:
                acid_match = False
                break

    if acid_match and acids_present:
        matrix.append({
            "check_item": "رقم القيد الجمركي المسبق (ACID Number 19-digits)",
            "status": "MATCH",
            "severity": "NONE",
            "system_value": sys_acid,
            "compared_values": {"coo": coo_acid, "inspection": insp_acid, "bl": bl_acid, "invoice": inv_acid},
            "details": f"مطابقة تامة لرقم الـ ACID ({first_acid}) عبر كافة المستندات.",
        })
    else:
        critical_discrepancies.append("اختلاف أو غياب في رقم الـ ACID بين المستندات الجمركية!")
        matrix.append({
            "check_item": "رقم القيد الجمركي المسبق (ACID Number 19-digits)",
            "status": "MISMATCH_CRITICAL",
            "severity": "BLOCKING",
            "system_value": sys_acid,
            "compared_values": {"coo": coo_acid, "inspection": insp_acid, "bl": bl_acid, "invoice": inv_acid},
            "details": "عدم تطابق رقم الـ ACID يمنع الإفراج الجمركي على منظومة نافذة.",
        })

    # 2. Importer / Consignee Cross-Check
    coo_imp = coo.get("importer_name") or ""
    insp_imp = insp.get("importer_name") or ""
    bl_imp = bl.get("consignee") or bl.get("importer_name") or ""
    m1, r1 = _fuzzy_match(sys_importer, coo_imp)
    m2, r2 = _fuzzy_match(sys_importer, insp_imp)
    m3, r3 = _fuzzy_match(sys_importer, bl_imp)
    imp_matched = (m1 or not coo_imp) and (m2 or not insp_imp) and (m3 or not bl_imp)

    matrix.append({
        "check_item": "اسم الشركة المستوردة (Importer / Consignee)",
        "status": "MATCH" if imp_matched else "MISMATCH_CRITICAL",
        "severity": "NONE" if imp_matched else "BLOCKING",
        "system_value": sys_importer,
        "compared_values": {"coo": coo_imp, "inspection": insp_imp, "bl": bl_imp},
        "details": "مطابقة اسم المستورد المصري" if imp_matched else "تطابق اسم المستورد مع السجل التجاري والبطاقة الاستيرادية مطلوب.",
    })
    if not imp_matched:
        critical_discrepancies.append("اختلاف اسم المستورد بين مسودة البوليصة والشهادات.")

    # 3. Exporter / Supplier Cross-Check
    coo_exp = coo.get("exporter_name") or ""
    insp_exp = insp.get("exporter_name") or ""
    bl_exp = bl.get("shipper") or bl.get("exporter_name") or ""
    e1, _ = _fuzzy_match(sys_supplier, coo_exp)
    e2, _ = _fuzzy_match(sys_supplier, insp_exp)
    e3, _ = _fuzzy_match(sys_supplier, bl_exp)
    exp_matched = (e1 or not coo_exp) and (e2 or not insp_exp) and (e3 or not bl_exp)

    matrix.append({
        "check_item": "المورد / المصدر الأجنبي (Exporter / Shipper)",
        "status": "MATCH" if exp_matched else "MISMATCH_CRITICAL",
        "severity": "NONE" if exp_matched else "BLOCKING",
        "system_value": sys_supplier,
        "compared_values": {"coo": coo_exp, "inspection": insp_exp, "bl": bl_exp},
        "details": "مطابقة تامة لبيانات المصنع والمورد الأجنبي" if exp_matched else "اختلاف في اسم أو عنوان المصدر الأجنبي.",
    })
    if not exp_matched:
        critical_discrepancies.append("اختلاف في بيانات المورد الأجنبي.")

    # 4. Gross Weight Cross-Check
    coo_wt = float(coo.get("gross_weight_kg") or 0.0)
    bl_wt = float(bl.get("total_gross_weight") or bl.get("gross_weight_kg") or 0.0)
    pkg_wt = float(pkg.get("total_gross_weight_kg") or 0.0)
    wts = [w for w in [coo_wt, bl_wt, pkg_wt] if w > 0]
    wt_matched = True
    if len(wts) >= 2:
        max_w = max(wts)
        min_w = min(wts)
        diff_pct = abs(max_w - min_w) / max_w * 100
        if diff_pct > 0.5:
            wt_matched = False
            warning_discrepancies.append(f"فرق في الوزن القائم بين المستندات بنسبة ({diff_pct:.2f}%).")

    sys_weight_val = getattr(imp_file, 'total_weight', None) or (pkg.get('total_gross_weight_kg') if pkg else None) or coo_wt or bl_wt or 0.0

    matrix.append({
        "check_item": "الوزن القائم الإجمالي (Total Gross Weight KG)",
        "status": "MATCH" if wt_matched else "MISMATCH_MINOR",
        "severity": "NONE" if wt_matched else "WARNING",
        "system_value": f"{sys_weight_val} KG",
        "compared_values": {"coo_weight": coo_wt, "bl_weight": bl_wt, "packing_weight": pkg_wt},
        "details": "الوزن متطابق ضمن نسبة السماح 0.5%" if wt_matched else "يرجى توحيد الوزن القائم في البوليصة وقائمة التعبئة لتفادي غرامات الجمارك.",
    })

    # 5. EUR.1 Revised Rules Exemption Check
    if "EUR" in str(coo.get("certificate_type", "")).upper():
        revised = coo.get("is_revised_rules", False) or "REVISED RULES" in str(coo.get("remarks", "")).upper()
        matrix.append({
            "check_item": "شرط القواعد المعدلة لليورو 1 (EUR.1 Box 7 REVISED RULES)",
            "status": "MATCH" if revised else "MISMATCH_CRITICAL",
            "severity": "NONE" if revised else "BLOCKING",
            "system_value": "REVISED RULES (إلزامي للإعفاء)",
            "compared_values": {"coo_remarks": coo.get("remarks", "")},
            "details": "مستوفٍ لقواعد المنشأ المعدلة وتطبيق ضريبة وارد 0%" if revised else "تحذير: غياب عبارة REVISED RULES قد يؤدي لرفض الإعفاء الجمركي.",
        })
        if not revised:
            critical_discrepancies.append("غياب عبارة REVISED RULES في الخانة 7 بشهادة EUR.1.")

    # 6. Inspection Draft Warning Check
    if insp.get("is_draft"):
        matrix.append({
            "check_item": "حالة شهادة الفحص (Inspection Cert Draft Status)",
            "status": "MISMATCH_MINOR",
            "severity": "WARNING",
            "system_value": "Original Final Certificate Required",
            "compared_values": {"status": "DRAFT", "warning": insp.get("draft_warning")},
            "details": "الشهادة مسودة DRAFT وتحتاج إلى تأكيد لإصدار النسخة النهائية المرقمة والموقعة.",
        })
        warning_discrepancies.append("شهادة الفحص والتفتيش الحالية مسودة وتحتاج إلى اعتماد نهائي.")

    is_safe = len(critical_discrepancies) == 0
    total_checks = len(matrix)
    passed_checks = len([m for m in matrix if m["status"] == "MATCH"])
    score = round((passed_checks / total_checks) * 100, 1) if total_checks > 0 else 100.0

    overall = "FULLY_MATCHED" if (is_safe and not warning_discrepancies) else ("ACCEPTED_WITH_WARNINGS" if is_safe else "DISCREPANCY_DETECTED")

    summary_ar = (
        f"اكتمل الفحص المتقاطع للمستندات بنسبة مطابقة {score}%. "
        + ("الشحنة مستوفية وجاهزة للربط والاعتماد الجمركي." if is_safe else f"يوجد {len(critical_discrepancies)} أخطاء حرجة تتطلب التصحيح الفوري قبل الشحن والإفراج.")
    )

    return ThreeWayCrossMatchResponse(
        import_file_id=request.import_file_id,
        import_file_code=imp_file.import_file_code or f"IMP-{request.import_file_id:04d}",
        overall_status=overall,
        match_score=score,
        is_safe_for_customs=is_safe,
        matrix=matrix,
        critical_discrepancies=critical_discrepancies,
        warning_discrepancies=warning_discrepancies,
        compliance_summary_ar=summary_ar,
    )


# --- 5. ACID & IMPORTER LEGAL DOCS EXPIRY & ETA + 30 DAYS SAFETY MARGIN ENGINE ---
def check_acid_and_company_docs_validity_service(db: Session, import_file_id: int) -> LegalDocsExpiryComplianceResponse:
    imp_file = db.query(ImportFile).filter(ImportFile.import_file_id == import_file_id).first()
    if not imp_file:
        raise HTTPException(status_code=404, detail="Import File not found")

    company = db.query(ImportCompany).filter(ImportCompany.company_id == imp_file.company_id).first() if imp_file.company_id else None
    booking = db.query(ShipmentBooking).filter(ShipmentBooking.import_file_id == import_file_id, ShipmentBooking.is_active == True).first()
    acid_session = db.query(AcidRegistrationSession).filter(AcidRegistrationSession.import_file_id == import_file_id, AcidRegistrationSession.is_active == True).first()

    today = date.today()

    # Determine Estimated Time of Arrival (ETA)
    eta_date: Optional[date] = None
    eta_source = "Estimated"

    if booking and booking.eta:
        eta_date = booking.eta.date() if hasattr(booking.eta, "date") else booking.eta
        eta_source = "Freight Booking (حجز الشحن الملاحي)"
    elif imp_file.required_eta:
        eta_date = imp_file.required_eta
        eta_source = "Import File (ملف الشحنة)"
    else:
        # Default safety assumption: 21 days from today
        from datetime import timedelta
        eta_date = today + timedelta(days=21)
        eta_source = "Estimated Default (+21 Days)"

    from datetime import timedelta
    safety_window_date = eta_date + timedelta(days=30)

    alerts: List[LegalDocAlertItem] = []
    has_critical_alerts = False

    # 1. ACID Expiry Check
    if acid_session and acid_session.expiry_date:
        acid_exp = acid_session.expiry_date
        days_rem = (acid_exp - today).days
        days_after_eta = (acid_exp - eta_date).days
        is_exp = days_rem <= 0
        is_breach = acid_exp <= safety_window_date
        
        if is_exp:
            status_str = "EXPIRED"
            msg = f"رقم الـ ACID ({acid_session.acid_number}) منتهي الصلاحية بتاريخ {acid_exp}. لا يمكن شحن البضائع أو الإفراج الجمركي!"
            has_critical_alerts = True
        elif is_breach:
            status_str = "CRITICAL_BREACH"
            msg = f"تحذير حرج: رقم الـ ACID ينتهي في {acid_exp} (أقل من 30 يوماً بعد تاريخ الوصول المتوقع {eta_date}). متبقي {days_after_eta} يوم فقط بعد الـ ETA."
            has_critical_alerts = True
        else:
            status_str = "VALID"
            msg = f"رقم الـ ACID ساري وصالح حتى {acid_exp} (صالح لمدة {days_after_eta} يوم بعد وصول الشحنة)."

        alerts.append(LegalDocAlertItem(
            doc_type="رقم الـ ACID للشحنة (Nafeza ACID)",
            doc_number=acid_session.acid_number,
            expiry_date=acid_exp,
            days_until_expiry=days_rem,
            days_after_eta=days_after_eta,
            is_expired=is_exp,
            is_critical_breach=is_breach,
            alert_message=msg,
            status=status_str,
        ))
    else:
        alerts.append(LegalDocAlertItem(
            doc_type="رقم الـ ACID للشحنة (Nafeza ACID)",
            doc_number=imp_file.acid_number or "لم يصدر بعد",
            expiry_date=None,
            days_until_expiry=0,
            days_after_eta=None,
            is_expired=False,
            is_critical_breach=True,
            alert_message="لم يتم استخراج أو تسجيل تاريخ صلاحية الـ ACID للشحنة بعد.",
            status="CRITICAL_BREACH",
        ))
        has_critical_alerts = True

    # 2. Importer Legal Documents Checks (Import Card, Commercial Registry, Tax Card)
    if company:
        docs_to_check = [
            ("البطاقة الاستيرادية (Import Card)", company.importer_id or "—", company.importer_id_expiry),
            ("السجل التجاري (Commercial Registry)", company.registration_number or "—", company.registration_expiry),
            ("البطاقة الضريبية / ملف القيمة المضافة (Tax Card)", company.vat_id or "—", company.vat_id_expiry),
        ]

        for doc_name, doc_num, exp_dt in docs_to_check:
            if exp_dt:
                days_rem = (exp_dt - today).days
                days_after_eta = (exp_dt - eta_date).days
                is_exp = days_rem <= 0
                is_breach = exp_dt <= safety_window_date

                if is_exp:
                    status_str = "EXPIRED"
                    msg = f"وثيقة {doc_name} منتهية بتاريخ {exp_dt}! يجب تجديدها فوراً وإلا ستتوقف إجراءات الإفراج الجمركي."
                    has_critical_alerts = True
                elif is_breach:
                    status_str = "CRITICAL_BREACH"
                    msg = f"تحذير حرج: {doc_name} تنتهي في {exp_dt} (قبل نافذة الـ 30 يوماً من وصول الشحنة في {eta_date})."
                    has_critical_alerts = True
                elif days_rem <= 60:
                    status_str = "EXPIRING_SOON"
                    msg = f"{doc_name} قاربت على الانتهاء في {exp_dt} (متبقي {days_rem} يوم)."
                else:
                    status_str = "VALID"
                    msg = f"{doc_name} سارية وصالحة حتى {exp_dt}."

                alerts.append(LegalDocAlertItem(
                    doc_type=doc_name,
                    doc_number=str(doc_num),
                    expiry_date=exp_dt,
                    days_until_expiry=days_rem,
                    days_after_eta=days_after_eta,
                    is_expired=is_exp,
                    is_critical_breach=is_breach,
                    alert_message=msg,
                    status=status_str,
                ))

    # Overall Compliance Calculation
    overall_status = "CRITICAL_ACTION_REQUIRED" if has_critical_alerts else "COMPLIANT"
    banner_text = None
    if has_critical_alerts:
        critical_msgs = [a.alert_message for a in alerts if a.is_critical_breach or a.is_expired]
        banner_text = "🚨 تحذير حرج لسلامة الشحنة والإفراج الجمركي (ETA + 30 Days Safety Margin):\n" + "\n".join(critical_msgs)

    return LegalDocsExpiryComplianceResponse(
        import_file_id=import_file_id,
        import_file_code=imp_file.import_file_code or f"IMP-{import_file_id}",
        company_name=company.importer_name if company else (imp_file.company_name or ""),
        eta_date=eta_date,
        eta_source=eta_source,
        safety_window_date=safety_window_date,
        has_critical_alerts=has_critical_alerts,
        total_alerts_count=len(alerts),
        alerts=alerts,
        overall_compliance_status=overall_status,
        persistent_banner_text=banner_text,
    )


# --- PHASE 6: COMMERCIAL INVOICE VS. BILL OF LADING CROSS-MATCHING SERVICE ---
def extract_and_match_invoice_bl_service(
    db: Session, request: InvoiceBLExtractAndMatchRequest
) -> InvoiceBLExtractAndMatchResponse:
    from modules.import_documentation.ai_document_parser import (
        extract_commercial_invoice_data,
        extract_packing_list_data,
        extract_draft_bl_with_ai,
        match_invoice_with_bl,
    )

    # 1. Parse Invoice
    invoice_data = dict(request.invoice_fields or {})
    if request.invoice_raw_text and request.invoice_raw_text.strip():
        extracted_inv = extract_commercial_invoice_data(request.invoice_raw_text)
        for k, v in extracted_inv.items():
            if k not in invoice_data or not invoice_data[k]:
                invoice_data[k] = v

    # 2. Parse Packing List (if provided as additional file / text)
    pl_data = dict(request.packing_list_fields or {})
    if request.packing_list_raw_text and request.packing_list_raw_text.strip():
        extracted_pl = extract_packing_list_data(request.packing_list_raw_text)
        for k, v in extracted_pl.items():
            if k not in pl_data or not pl_data[k]:
                pl_data[k] = v

    # Merge Packing List parameters into invoice_data to ensure complete cross-matching
    if pl_data:
        if not invoice_data.get("total_gross_weight_kg") and pl_data.get("total_gross_weight_kg"):
            invoice_data["total_gross_weight_kg"] = pl_data["total_gross_weight_kg"]
        if not invoice_data.get("total_net_weight_kg") and pl_data.get("total_net_weight_kg"):
            invoice_data["total_net_weight_kg"] = pl_data["total_net_weight_kg"]
        if not invoice_data.get("total_packages") and pl_data.get("total_packages"):
            invoice_data["total_packages"] = pl_data["total_packages"]
            invoice_data["qty_pkg"] = pl_data["total_packages"]
        if not invoice_data.get("cbm") and pl_data.get("total_cbm"):
            invoice_data["cbm"] = pl_data["total_cbm"]
        if not invoice_data.get("containers") and pl_data.get("containers"):
            invoice_data["containers"] = pl_data["containers"]
        if not invoice_data.get("acid_number") and pl_data.get("acid_number"):
            invoice_data["acid_number"] = pl_data["acid_number"]
        if not invoice_data.get("importer_tax_id") and pl_data.get("importer_tax_id"):
            invoice_data["importer_tax_id"] = pl_data["importer_tax_id"]
        if not invoice_data.get("shipper") and pl_data.get("shipper"):
            invoice_data["shipper"] = pl_data["shipper"]
        if not invoice_data.get("consignee") and pl_data.get("consignee"):
            invoice_data["consignee"] = pl_data["consignee"]


    # 3. Parse Bill of Lading
    bl_data = dict(request.bl_fields or {})
    if request.bl_raw_text and request.bl_raw_text.strip():
        extracted_bl = extract_draft_bl_with_ai(request.bl_raw_text)
        for k, v in extracted_bl.items():
            if k not in bl_data or not bl_data[k]:
                bl_data[k] = v

    # 4. System snapshot if import_file_id is provided
    sys_data = None
    imp_file_code = None
    if request.import_file_id:
        sys_data = _build_system_bl_snapshot(db, request.import_file_id)
        imp_file = db.query(ImportFile).filter(ImportFile.import_file_id == request.import_file_id).first()
        if imp_file:
            imp_file_code = imp_file.import_file_code

    # 5. Execute Smart Cross-Matching
    match_result = match_invoice_with_bl(invoice_data, bl_data, sys_data)

    return InvoiceBLExtractAndMatchResponse(
        import_file_id=request.import_file_id,
        import_file_code=imp_file_code,
        overall_status=match_result["overall_status"],
        is_safe_for_certification=match_result["is_safe_for_certification"],
        match_score_percentage=match_result["match_score_percentage"],
        critical_discrepancies_count=match_result["critical_discrepancies_count"],
        warning_discrepancies_count=match_result["warning_discrepancies_count"],
        comparison_matrix=match_result["comparison_matrix"],
        correction_letter=match_result["correction_letter"],
        invoice_data=match_result["invoice_data"],
        bl_data=match_result["bl_data"],
        packing_list_data=pl_data or None,
    )



def sync_certified_invoice_bl_to_file_service(
    db: Session, request: InvoiceBLSyncRequest
) -> dict:
    imp_file = db.query(ImportFile).filter(ImportFile.import_file_id == request.import_file_id).first()
    if not imp_file:
        raise HTTPException(status_code=404, detail="Import File not found")

    inv = request.invoice_data or {}
    bl = request.bl_data or {}

    # Sync Invoice fields
    if inv.get("invoice_number"):
        imp_file.pi_number = str(inv["invoice_number"])
    if inv.get("total_amount"):
        try:
            imp_file.total_amount = float(inv["total_amount"])
        except (ValueError, TypeError):
            pass
    if inv.get("currency"):
        imp_file.currency = str(inv["currency"])
    if inv.get("incoterm"):
        imp_file.incoterm = str(inv["incoterm"])

    # Sync B/L fields
    if bl.get("draft_bl_number"):
        imp_file.bl_number = str(bl["draft_bl_number"])
    if bl.get("booking_no"):
        imp_file.booking_no = str(bl["booking_no"])
    if bl.get("vessel_name"):
        imp_file.vessel_name = str(bl["vessel_name"])

    # Sync to Cargo Shipping Record if requested
    if request.sync_to_shipping:
        cargo_shp = db.query(CargoShippingRecord).filter(
            CargoShippingRecord.import_file_id == request.import_file_id,
            CargoShippingRecord.is_active == True
        ).first()
        if cargo_shp:
            if bl.get("draft_bl_number"):
                cargo_shp.bl_number = str(bl["draft_bl_number"])
            if bl.get("vessel_name"):
                cargo_shp.vessel_name = str(bl["vessel_name"])
            if bl.get("containers"):
                cargo_shp.containers_loading_data = bl["containers"]
            if bl.get("total_gross_weight_kg"):
                cargo_shp.total_gross_weight_kg = float(bl["total_gross_weight_kg"])

    db.commit()
    db.refresh(imp_file)

    return {
        "status": "success",
        "message": f"تمت مطابقة واعتماد الفاتورة ({inv.get('invoice_number', 'N/A')}) والبوليصة ({bl.get('draft_bl_number', 'N/A')}) ومزامنة بياناتهما مع ملف الشحنة بنجاح.",
        "import_file_id": imp_file.import_file_id,
        "import_file_code": imp_file.import_file_code,
        "synced_invoice_number": imp_file.pi_number,
        "synced_bl_number": imp_file.bl_number,
        "total_amount": imp_file.total_amount,
        "currency": imp_file.currency,
    }


def extract_and_compare_po_documents_service(
    db: Session, request: POExtractAndCompareRequest
) -> POExtractAndCompareResponse:
    from modules.import_documentation.ai_document_parser import (
        extract_commercial_invoice_data,
        extract_packing_list_data,
        reconcile_po_documents_with_system,
    )
    from modules.import_files.model import ImportFile

    # 1. Extract Invoice Data
    inv_data = dict(request.invoice_data or {})
    if request.invoice_raw_text and request.invoice_raw_text.strip():
        extracted_inv = extract_commercial_invoice_data(request.invoice_raw_text)
        for k, v in extracted_inv.items():
            if k not in inv_data or not inv_data[k]:
                inv_data[k] = v

    # 2. Extract Packing List Data
    pl_data = dict(request.packing_data or {})
    if request.packing_list_raw_text and request.packing_list_raw_text.strip():
        extracted_pl = extract_packing_list_data(request.packing_list_raw_text)
        for k, v in extracted_pl.items():
            if k not in pl_data or not pl_data[k]:
                pl_data[k] = v

    # 3. Fetch System PO Items and File Metadata if import_file_id is provided
    system_items = list(request.system_items or [])
    file_metadata = {}
    if request.import_file_id:
        imp_file = db.query(ImportFile).filter(ImportFile.import_file_id == request.import_file_id).first()
        if imp_file:
            tax_no = None
            if imp_file.company and hasattr(imp_file.company, 'tax_card_number'):
                tax_no = imp_file.company.tax_card_number
            file_metadata = {
                "import_file_id": imp_file.import_file_id,
                "import_file_code": imp_file.import_file_code,
                "acid_number": imp_file.acid_number,
                "importer_tax_id": tax_no,
                "total_amount": imp_file.total_amount or 0.0,
                "currency": imp_file.currency or "EUR",
                "total_packages": imp_file.total_packages or 0,
                "total_gross_weight_kg": imp_file.total_gross_weight_kg or 0.0,
            }
            if not system_items:
                from modules.purchase_orders.model import PurchaseOrder
                po_records = db.query(PurchaseOrder).filter(PurchaseOrder.import_file_id == request.import_file_id).all()
                for po in po_records:
                    for itm in po.items:
                        system_items.append({
                            "po_item_id": itm.item_id,
                            "item_code": itm.item_code or str(itm.item_id),
                            "description": itm.description_ar or itm.description_en or "بند أمر الشراء",
                            "hs_code": itm.hs_code or "",
                            "package_type": "Carton",
                            "initial_quantity": itm.quantity,
                            "initial_unit_price": itm.unit_price,
                            "initial_packages_count": 1.0,
                            "initial_net_weight_kg": itm.net_weight_kg or 0.0,
                            "initial_gross_weight_kg": itm.gross_weight_kg or 0.0,
                            "initial_cbm": itm.total_cbm or 0.0,
                        })

    if not system_items:
        inv_itms = inv_data.get("items", [])
        for idx, itm in enumerate(inv_itms, 1):
            system_items.append({
                "po_item_id": idx,
                "item_code": itm.get("item_code", f"ITEM-{idx}"),
                "description": itm.get("description", f"بند {idx}"),
                "hs_code": itm.get("hs_code", ""),
                "package_type": "Package",
                "initial_quantity": itm.get("quantity", 1.0),
                "initial_unit_price": itm.get("unit_price", 0.0),
                "initial_packages_count": 1.0,
                "initial_net_weight_kg": 0.0,
                "initial_gross_weight_kg": 0.0,
                "initial_cbm": 0.0,
            })

    # 4. Perform 3-Way Reconciliation
    reconciled = reconcile_po_documents_with_system(inv_data, pl_data, system_items, file_metadata)

    return POExtractAndCompareResponse(
        import_file_id=request.import_file_id,
        overall_status=reconciled["overall_status"],
        is_safe_for_certification=reconciled["is_safe_for_certification"],
        critical_discrepancies_count=reconciled["critical_discrepancies_count"],
        warning_discrepancies_count=reconciled["warning_discrepancies_count"],
        header_discrepancies=reconciled["header_discrepancies"],
        reconciled_invoice_items=reconciled["reconciled_invoice_items"],
        reconciled_packing_items=reconciled["reconciled_packing_items"],
        extracted_invoice_data=reconciled["extracted_invoice_data"],
        extracted_packing_data=reconciled["extracted_packing_data"],
    )


# --- PO & PACKING RECONCILIATION SESSIONS SERVICES (BP-016) ---
def enrich_po_reconciliation_session_response(
    db: Session, session: POPackingReconciliationSession
) -> POReconciliationSessionResponse:
    import_file_code = None
    importer_name = None

    if session.import_file_id:
        file = db.query(ImportFile).filter(ImportFile.import_file_id == session.import_file_id).first()
        if file:
            import_file_code = file.import_file_code
            importer_name = file.company_name or (file.company.importer_name if file.company else None)


    resp_dict = {
        "session_id": session.session_id,
        "session_code": session.session_code,
        "import_file_id": session.import_file_id,
        "import_file_code": import_file_code,
        "importer_name": importer_name,
        "final_invoice_number": session.final_invoice_number,
        "final_packing_list_number": session.final_packing_list_number,
        "acid_number": session.acid_number,
        "shipper_name": session.shipper_name,
        "total_invoice_amount": session.total_invoice_amount,
        "currency": session.currency,
        "total_packages": session.total_packages,
        "total_net_weight_kg": session.total_net_weight_kg,
        "total_gross_weight_kg": session.total_gross_weight_kg,
        "total_cbm": session.total_cbm,
        "overall_status": session.overall_status,
        "is_safe_for_certification": session.is_safe_for_certification,
        "critical_discrepancies_count": session.critical_discrepancies_count,
        "warning_discrepancies_count": session.warning_discrepancies_count,
        "header_discrepancies": session.header_discrepancies,
        "reconciled_invoice_items": session.reconciled_invoice_items,
        "reconciled_packing_items": session.reconciled_packing_items,
        "extracted_invoice_data": session.extracted_invoice_data,
        "extracted_packing_data": session.extracted_packing_data,
        "notes": session.notes,
        "certified_by": session.certified_by,
        "is_active": session.is_active,
        "created_at": session.created_at,
        "updated_at": session.updated_at,
    }
    return POReconciliationSessionResponse(**resp_dict)


def create_po_reconciliation_session_service(
    db: Session, schema: POReconciliationSessionCreate
) -> POReconciliationSessionResponse:
    # 1. Validate import file exists
    file = db.query(ImportFile).filter(ImportFile.import_file_id == schema.import_file_id).first()
    if not file:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"الملف الاستيرادي رقم {schema.import_file_id} غير موجود بالمنظومة.",
        )

    # 2. Strict Rule: Prevent duplicate session for the same import file!
    existing = repo.get_po_reconciliation_session_by_file_id(db, schema.import_file_id)
    if existing:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"يوجد بالفعل جلسة مطابقة محفوظة لهذا الملف الاستيرادي (رمز الجلسة: {existing.session_code}). "
                f"لا يُسمح بإنشاء أكثر من جلسة حفظ لنفس الملف الاستيرادي. يمكنك تحديث الجلسة الحالية أو استعراضها من سجل الجلسات."
            ),
        )

    session = repo.create_po_reconciliation_session(db, schema)
    return enrich_po_reconciliation_session_response(db, session)


def get_po_reconciliation_sessions_service(
    db: Session,
    import_file_id: int | None = None,
    overall_status: str | None = None,
    search: str | None = None,
) -> list[POReconciliationSessionResponse]:
    sessions = repo.get_po_reconciliation_sessions(
        db, import_file_id=import_file_id, overall_status=overall_status, search=search
    )
    return [enrich_po_reconciliation_session_response(db, s) for s in sessions]


def get_po_reconciliation_session_by_id_service(
    db: Session, session_id: int
) -> POReconciliationSessionResponse:
    session = repo.get_po_reconciliation_session_by_id(db, session_id)
    if not session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"جلسة المطابقة رقم {session_id} غير موجودة.",
        )
    return enrich_po_reconciliation_session_response(db, session)


def update_po_reconciliation_session_service(
    db: Session, session_id: int, schema: POReconciliationSessionUpdate
) -> POReconciliationSessionResponse:
    session = repo.get_po_reconciliation_session_by_id(db, session_id)
    if not session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"جلسة المطابقة رقم {session_id} غير موجودة للتحديث.",
        )
    updated = repo.update_po_reconciliation_session(db, session_id, schema)
    return enrich_po_reconciliation_session_response(db, updated)


def delete_po_reconciliation_session_service(db: Session, session_id: int) -> dict:
    ok = repo.delete_po_reconciliation_session(db, session_id)
    if not ok:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"جلسة المطابقة رقم {session_id} غير موجودة للحذف.",
        )
    return {"message": f"تم حذف جلسة المطابقة رقم {session_id} بنجاح", "deleted": True}





def get_all_acid_sessions_service(db: Session, include_inactive: bool = False, search: str = None, import_file_id: int = None, status: str = None):
    return repo.get_all_acid_sessions(db, include_inactive=include_inactive, search=search, import_file_id=import_file_id, status=status)

def get_acid_session_by_id_service(db: Session, acid_id: int):
    return repo.get_acid_session_by_id(db, acid_id)

def soft_delete_acid_session_service(db: Session, acid_id: int):
    return repo.soft_delete_acid_session(db, acid_id)

def get_all_banking_documents_service(db: Session, import_file_id: int = None):
    return repo.get_all_banking_documents(db, import_file_id=import_file_id)

def get_banking_document_by_id_service(db: Session, bank_doc_id: int):
    return repo.get_banking_document_by_id(db, bank_doc_id)

def get_all_shipment_documents_service(db: Session, import_file_id: int = None):
    return repo.get_all_shipment_documents(db, import_file_id=import_file_id)

def get_draft_bl_reviews_service(db: Session, include_inactive: bool = False, import_file_id: int = None, status: str = None, search: str = None):
    return repo.get_draft_bl_reviews(db, include_inactive=include_inactive, import_file_id=import_file_id, status=status, search=search)

def get_draft_bl_review_by_id_service(db: Session, review_id: int, include_inactive: bool = True):
    return repo.get_draft_bl_review_by_id(db, review_id, include_inactive=include_inactive)

def get_coo_reviews_service(db: Session, include_inactive: bool = False, import_file_id: int = None, status: str = None, search: str = None):
    return repo.get_coo_reviews(db, include_inactive=include_inactive, import_file_id=import_file_id, status=status, search=search)

def get_inspection_reviews_service(db: Session, include_inactive: bool = False, import_file_id: int = None, status: str = None, search: str = None):
    return repo.get_inspection_reviews(db, include_inactive=include_inactive, import_file_id=import_file_id, status=status, search=search)


def get_central_archive_service(db: Session, import_file_id: int) -> CentralArchiveResponse:
    from modules.import_requirements.model import ImportRequirementAssessment
    from modules.import_files.model import ImportFile
    from modules.import_companies.model import ImportCompany
    from modules.suppliers.model import Supplier
    from modules.cargo_shipping.model import ShipmentBooking

    imp_file = db.query(ImportFile).filter(ImportFile.import_file_id == import_file_id).first()
    if not imp_file:
        raise HTTPException(status_code=404, detail="Import File not found")

    company = db.query(ImportCompany).filter(ImportCompany.company_id == imp_file.company_id).first() if imp_file.company_id else None
    supplier = db.query(Supplier).filter(Supplier.supplier_id == imp_file.supplier_id).first() if imp_file.supplier_id else None
    booking = db.query(ShipmentBooking).filter(
        ShipmentBooking.import_file_id == import_file_id,
        ShipmentBooking.is_active == True
    ).first()
    acid_session = db.query(AcidRegistrationSession).filter(
        AcidRegistrationSession.import_file_id == import_file_id,
        AcidRegistrationSession.is_active == True
    ).first()
    po_reconciliation = db.query(POPackingReconciliationSession).filter(
        POPackingReconciliationSession.import_file_id == import_file_id,
        POPackingReconciliationSession.is_active == True
    ).first()
    bl_review = db.query(DraftBLReviewSession).filter(
        DraftBLReviewSession.import_file_id == import_file_id,
        DraftBLReviewSession.is_active == True
    ).order_by(DraftBLReviewSession.bl_review_id.desc()).first()
    coo_review = db.query(CertificateOfOriginReviewSession).filter(
        CertificateOfOriginReviewSession.import_file_id == import_file_id,
        CertificateOfOriginReviewSession.is_active == True
    ).order_by(CertificateOfOriginReviewSession.coo_review_id.desc()).first()
    inspection_review = db.query(InspectionCertificateReviewSession).filter(
        InspectionCertificateReviewSession.import_file_id == import_file_id,
        InspectionCertificateReviewSession.is_active == True
    ).order_by(InspectionCertificateReviewSession.inspection_review_id.desc()).first()

    # Query Import Requirement Assessment (BP-011)
    assessment = db.query(ImportRequirementAssessment).filter(
        ImportRequirementAssessment.import_file_id == import_file_id,
        ImportRequirementAssessment.is_active == True
    ).order_by(ImportRequirementAssessment.assessment_id.desc()).first()
    if not assessment and imp_file.supplier_id:
        assessment = db.query(ImportRequirementAssessment).filter(
            ImportRequirementAssessment.supplier_id == imp_file.supplier_id,
            ImportRequirementAssessment.is_active == True
        ).order_by(ImportRequirementAssessment.assessment_id.desc()).first()

    # Determine Country of Origin
    origin_country = (
        (assessment.country_of_origin if assessment and assessment.country_of_origin else None)
        or (coo_review.country_of_origin if coo_review and coo_review.country_of_origin else None)
        or (supplier.foreign_exporter_country if supplier and supplier.foreign_exporter_country else None)
        or (getattr(imp_file, 'country_of_origin', None))
        or "European Union / International"
    )

    eu_countries = [
        'lithuania', 'germany', 'italy', 'france', 'spain', 'poland', 'belgium',
        'netherlands', 'austria', 'czech', 'sweden', 'denmark', 'finland', 'greece',
        'portugal', 'hungary', 'ireland', 'romania', 'bulgaria', 'slovakia', 'slovenia',
        'croatia', 'latvia', 'estonia', 'cyprus', 'malta', 'luxembourg', 'eu', 'europe', 'european'
    ]
    is_eu_origin = any(c in origin_country.lower() for c in eu_countries) or (coo_review and coo_review.certificate_type == "EUR.1")

    # Evaluate COO / EUR.1 Requirements (Rule 2: Conditional by Tariff/Origin)
    tariff_exemption_alert = None
    if assessment and assessment.coo_required:
        coo_is_mandatory = True
        coo_is_waived = False
        if is_eu_origin or assessment.coo_type == "EUR.1":
            coo_legal_note = "مطلوبة شهادة EUR.1 مدوناً بها عبارة 'REVISED RULES' بالخانة 7 لتطبيق إعفاء 0% ضريبة وارد."
            tariff_exemption_alert = "⚠️ فرصة إعفاء جمركي: الشحنة مؤهلة لإعفاء كامل (0% ضريبة وارد) بموجب الشراكة المصرية الأوروبية، بشرط تقديم شهادة EUR.1 مدوناً بها عبارة 'REVISED RULES' بالخانة 7."
        else:
            coo_legal_note = "مطلوبة شهادة منشأ معتمدة لإثبات المنشأ الجمركي ومطابقة التعرفة."
    elif assessment and not assessment.coo_required:
        coo_is_mandatory = False
        coo_is_waived = True
        coo_legal_note = "غير مطلوبة / معفاة قانونياً وفق تقييم متطلبات الاستيراد (BP-011)."
    else:
        # Default heuristic based on origin
        if is_eu_origin:
            coo_is_mandatory = True
            coo_is_waived = False
            coo_legal_note = "مطلوبة شهادة EUR.1 بعبارة 'REVISED RULES' للاستفادة من الإعفاء الجمركي 0%."
            tariff_exemption_alert = "⚠️ فرصة إعفاء جمركي: الشحنة مؤهلة لإعفاء كامل (0% ضريبة وارد) بموجب الشراكة المصرية الأوروبية، بشرط تقديم شهادة EUR.1 مدوناً بها عبارة 'REVISED RULES' بالخانة 7."
        elif "china" in origin_country.lower() or (coo_review and "ccpit" in str(coo_review.certificate_type).lower()):
            coo_is_mandatory = True
            coo_is_waived = False
            coo_legal_note = "مطلوبة شهادة منشأ الصين CCPIT إلكترونية مع باركود ورابط التحقق."
        else:
            coo_is_mandatory = False
            coo_is_waived = True
            coo_legal_note = "شهادة المنشأ غير ملزمة / معفاة لهذا البند الجمركي والمنشأ."

    # Evaluate Inspection Certificate (VoC / COC) Requirements (Rule 2: Conditional by GOEIC/NFSA)
    goeic_inspection_alert = None
    if assessment and assessment.inspection_required:
        insp_is_mandatory = True
        insp_is_waived = False
        insp_agency = assessment.inspection_body or "COTECNA / TÜV / SGS"
        insp_legal_note = f"خاضع لفحص الهيئة العامة للرقابة على الصادرات والواردات (GOEIC) وتفتيش ما قبل الشحن عبر ({insp_agency})."
        goeic_inspection_alert = f"⚠️ تنبيه فحص نوعي (GOEIC): البند الجمركي خاضع لفحص ما قبل الشحن من شركة معاينة معتمدة ({insp_agency}) مع ضرورة تدقيق مسودة الـ 48 ساعة."
    elif assessment and not assessment.inspection_required:
        insp_is_mandatory = False
        insp_is_waived = True
        insp_legal_note = "الصنف غير خاضع لفحص الصادرات والواردات (GOEIC) ولا يتطلب شهادة فحص ما قبل الشحن."
    else:
        # If an inspection review session exists, consider it active; otherwise waived
        if inspection_review is not None:
            insp_is_mandatory = True
            insp_is_waived = False
            insp_legal_note = f"تم إدراج فحص نوعي ما قبل الشحن عبر ({inspection_review.inspection_agency})."
        else:
            insp_is_mandatory = False
            insp_is_waived = True
            insp_legal_note = "غير خاضع لفحص ما قبل الشحن بموجب البند الجمركي."

    # Evaluate Decree 43 / White List
    decree_43_alert = None
    if assessment and assessment.decree_43_applicable:
        if not assessment.white_list_verified:
            supp_name = assessment.supplier_name or (supplier.company_name if supplier else (imp_file.supplier_name or "المورد"))
            decree_43_alert = f"⚠️ تنبيه قرار 43: المصنع ({supp_name}) خاضع للقرار 43/2016 ويلزم التحقق من قيده بالقائمة البيضاء للهيئة برقم: {assessment.factory_registration_no or 'غير مسجل'}"

    all_discrepancies = []
    total_critical = 0
    total_warning = 0

    # 1. Final Commercial Invoice (100% Mandatory Core)
    inv_available = po_reconciliation is not None and bool(po_reconciliation.final_invoice_number)
    inv_discrepancies = []
    if po_reconciliation and po_reconciliation.header_discrepancies:
        for d in po_reconciliation.header_discrepancies:
            if "invoice" in str(d.get("field", "")).lower() or "price" in str(d.get("field", "")).lower() or "amount" in str(d.get("field", "")).lower():
                sev = "CRITICAL" if d.get("severity") in ["CRITICAL", "BLOCKING"] else "WARNING"
                item = {
                    "document": "الفاتورة التجارية النهائية (Commercial Invoice)",
                    "field": d.get("field", "Invoice Field"),
                    "severity": sev,
                    "issue": d.get("issue", d.get("description", "عدم تطابق في الفاتورة")),
                    "rectification": d.get("rectification", "تعديل الفاتورة التجارية لتطابق أمر الشراء"),
                }
                inv_discrepancies.append(item)
                all_discrepancies.append(item)
                if sev == "CRITICAL": total_critical += 1
                else: total_warning += 1

    final_invoice = CentralArchiveDocumentSummary(
        document_type="FINAL_COMMERCIAL_INVOICE",
        title_ar="الفاتورة التجارية النهائية المعتمدة",
        is_available=inv_available,
        is_mandatory=True,
        is_waived=False,
        legal_requirement_note="مستند إلزامي حتمي لكافة الشحنات لإثبات القيمة الجمركية (CIF/FOB) والوعاء الضريبي.",
        status="APPROVED" if (inv_available and po_reconciliation.is_safe_for_certification) else ("REVIEW_PENDING" if inv_available else "NOT_STARTED"),
        document_reference=(po_reconciliation.final_invoice_number if po_reconciliation else None) or imp_file.pi_number,
        details={
            "invoice_number": (po_reconciliation.final_invoice_number if po_reconciliation else None) or imp_file.pi_number,
            "total_amount": float(po_reconciliation.total_invoice_amount) if po_reconciliation else float(getattr(imp_file, 'estimated_cost', 0.0) or 0.0),
            "currency": (po_reconciliation.currency if po_reconciliation else None) or getattr(imp_file, 'estimated_cost_currency', 'EUR'),
            "shipper_name": (po_reconciliation.shipper_name if po_reconciliation else None) or (supplier.company_name if supplier else imp_file.supplier_name),
            "reconciled_items_count": len(po_reconciliation.reconciled_invoice_items or []) if po_reconciliation else 0,
            "certified_by": po_reconciliation.certified_by if po_reconciliation else None,
        },
        discrepancies=inv_discrepancies,
        last_updated=po_reconciliation.updated_at if po_reconciliation else None
    )

    # 2. Final Packing List (100% Mandatory Core)
    pkg_available = po_reconciliation is not None and bool(po_reconciliation.final_packing_list_number or po_reconciliation.total_packages)
    pkg_discrepancies = []
    if po_reconciliation and po_reconciliation.header_discrepancies:
        for d in po_reconciliation.header_discrepancies:
            if "weight" in str(d.get("field", "")).lower() or "package" in str(d.get("field", "")).lower() or "cbm" in str(d.get("field", "")).lower():
                sev = "CRITICAL" if d.get("severity") in ["CRITICAL", "BLOCKING"] else "WARNING"
                item = {
                    "document": "قائمة التعبئة النهائية (Packing List)",
                    "field": d.get("field", "Packing Field"),
                    "severity": sev,
                    "issue": d.get("issue", d.get("description", "عدم تطابق في التعبئة أو الوزن")),
                    "rectification": d.get("rectification", "تعديل قائمة التعبئة لتطابق الأوزان الفعلية"),
                }
                pkg_discrepancies.append(item)
                all_discrepancies.append(item)
                if sev == "CRITICAL": total_critical += 1
                else: total_warning += 1

    final_packing_list = CentralArchiveDocumentSummary(
        document_type="FINAL_PACKING_LIST",
        title_ar="قائمة التعبئة النهائية المعتمدة",
        is_available=pkg_available,
        is_mandatory=True,
        is_waived=False,
        legal_requirement_note="مستند إلزامي حتمي لتحديد الأوزان الصافية والقائمة وعدد الطرود وأرقام الحاويات للجمارك.",
        status="APPROVED" if (pkg_available and po_reconciliation.is_safe_for_certification) else ("REVIEW_PENDING" if pkg_available else "NOT_STARTED"),
        document_reference=(po_reconciliation.final_packing_list_number if po_reconciliation else None) or f"PL-{imp_file.import_file_code}",
        details={
            "packing_list_number": (po_reconciliation.final_packing_list_number if po_reconciliation else None) or f"PL-{imp_file.import_file_code}",
            "total_packages": int(po_reconciliation.total_packages) if po_reconciliation else 0,
            "total_gross_weight_kg": float(po_reconciliation.total_gross_weight_kg) if po_reconciliation else 0.0,
            "total_net_weight_kg": float(po_reconciliation.total_net_weight_kg) if po_reconciliation else 0.0,
            "total_cbm": float(po_reconciliation.total_cbm) if po_reconciliation else 0.0,
        },
        discrepancies=pkg_discrepancies,
        last_updated=po_reconciliation.updated_at if po_reconciliation else None
    )

    # 3. Draft Bill of Lading (100% Mandatory Core)
    bl_available = bl_review is not None
    bl_discrepancies = []
    if bl_review and bl_review.comparison_matrix:
        for row in bl_review.comparison_matrix:
            if row.get("status") in ["MISMATCH", "MISMATCH_CRITICAL", "MISMATCH_MINOR", "DIFF"]:
                sev = "CRITICAL" if row.get("severity") in ["BLOCKING", "CRITICAL"] or row.get("status") == "MISMATCH_CRITICAL" else "WARNING"
                item = {
                    "document": "مسودة بوليصة الشحن (Draft B/L)",
                    "field": row.get("field") or row.get("check_item", "B/L Field"),
                    "severity": sev,
                    "issue": f"القيمة في البوليصة: '{row.get('bl_val') or row.get('compared_value')}' مقابل النظام: '{row.get('sys_val') or row.get('system_value')}'",
                    "rectification": row.get("correction_instruction") or row.get("details") or "طلب تعديل درافت البوليصة من الخط الملاحي",
                }
                bl_discrepancies.append(item)
                all_discrepancies.append(item)
                if sev == "CRITICAL": total_critical += 1
                else: total_warning += 1

    draft_bl = CentralArchiveDocumentSummary(
        document_type="DRAFT_BL",
        title_ar="مسودة بوليصة الشحن (Draft B/L)",
        is_available=bl_available,
        is_mandatory=True,
        is_waived=False,
        legal_requirement_note="مستند إلزامي حتمي يثبت الشحن والناقل البحري ورقم القيد الجمركي ACID.",
        status="APPROVED" if (bl_available and bl_review.status in ["Approved", "APPROVED", "FINAL"]) else ("MODIFICATIONS_REQUESTED" if bl_discrepancies else ("REVIEW_PENDING" if bl_available else "NOT_STARTED")),
        document_reference=(bl_review.draft_bl_number if bl_review else None) or (booking.booking_reference if booking else None),
        details={
            "draft_bl_number": (bl_review.draft_bl_number if bl_review else None) or "DRAFT-BL",
            "shipping_line": (bl_review.shipping_line if bl_review else None) or (booking.shipping_line if booking else "MSC / Maersk"),
            "vessel_name": (bl_review.vessel_name if bl_review else None) or (booking.vessel_name if booking else None),
            "voyage_number": bl_review.voyage_number if bl_review else None,
            "container_summary": bl_review.container_summary if bl_review else None,
            "gross_weight_kg": bl_review.gross_weight if (bl_review and hasattr(bl_review, 'gross_weight')) else (booking.gross_weight if booking else 0.0),
        },
        discrepancies=bl_discrepancies,
        raw_content=bl_review.raw_input_text if (bl_review and hasattr(bl_review, 'raw_input_text')) else None,
        last_updated=bl_review.updated_at if bl_review else None
    )

    # 4. Certificate of Origin / EUR.1 (Conditional)
    coo_available = coo_review is not None
    coo_discrepancies = []
    if coo_review and coo_review.comparison_matrix:
        for row in coo_review.comparison_matrix:
            if row.get("status") in ["MISMATCH", "MISMATCH_CRITICAL", "MISMATCH_MINOR", "DIFF"]:
                sev = "CRITICAL" if row.get("severity") in ["BLOCKING", "CRITICAL"] or row.get("status") == "MISMATCH_CRITICAL" else "WARNING"
                item = {
                    "document": f"شهادة المنشأ ({coo_review.certificate_type})",
                    "field": row.get("field") or row.get("check_item", "COO Field"),
                    "severity": sev,
                    "issue": f"المطابقة: {row.get('details') or row.get('issue', 'فرق في بيانات المنشأ')}",
                    "rectification": row.get("rectification", "طلب تعديل درافت شهادة المنشأ من المصدر أو الغرفة التجارية"),
                }
                coo_discrepancies.append(item)
                all_discrepancies.append(item)
                if sev == "CRITICAL": total_critical += 1
                else: total_warning += 1

    coo_status = "WAIVED" if coo_is_waived else (
        "APPROVED" if (coo_available and coo_review.status in ["Approved", "APPROVED", "Verified"]) else ("MODIFICATIONS_REQUESTED" if coo_discrepancies else ("REVIEW_PENDING" if coo_available else "NOT_STARTED"))
    )

    certificate_of_origin = CentralArchiveDocumentSummary(
        document_type="CERTIFICATE_OF_ORIGIN",
        title_ar="شهادة المنشأ / يورو 1 (COO / EUR.1)",
        is_available=coo_available,
        is_mandatory=coo_is_mandatory,
        is_waived=coo_is_waived,
        waive_reason="غير مطلوبة قانونياً / معفاة وفق تقييم البند الجمركي (BP-011)" if coo_is_waived else None,
        legal_requirement_note=coo_legal_note,
        status=coo_status,
        document_reference=coo_review.certificate_number if coo_review else None,
        details={
            "certificate_type": (coo_review.certificate_type if coo_review else ("EUR.1" if is_eu_origin else "Standard COO")),
            "certificate_number": coo_review.certificate_number if coo_review else "DRAFT-COO",
            "country_of_origin": origin_country,
            "exporter_name": (coo_review.exporter_name if coo_review else (supplier.company_name if supplier else imp_file.supplier_name)),
            "importer_name": (coo_review.importer_name if coo_review else (company.importer_name if company else imp_file.company_name)),
            "is_preferential_agreement": is_eu_origin,
            "tariff_exemption_eligible": is_eu_origin,
        },
        discrepancies=coo_discrepancies,
        raw_content=coo_review.raw_input_text if coo_review else None,
        last_updated=coo_review.updated_at if coo_review else None
    )

    # 5. Inspection Certificate / VoC (Conditional)
    insp_available = inspection_review is not None
    insp_discrepancies = []
    if inspection_review and inspection_review.comparison_matrix:
        for row in inspection_review.comparison_matrix:
            if row.get("status") in ["MISMATCH", "MISMATCH_CRITICAL", "MISMATCH_MINOR", "DIFF"]:
                sev = "CRITICAL" if row.get("severity") in ["BLOCKING", "CRITICAL"] or row.get("status") == "MISMATCH_CRITICAL" else "WARNING"
                item = {
                    "document": f"شهادة الفحص ({inspection_review.inspection_agency})",
                    "field": row.get("field") or row.get("check_item", "Inspection Field"),
                    "severity": sev,
                    "issue": f"المطابقة: {row.get('details') or row.get('issue', 'فرق في بيانات الفحص')}",
                    "rectification": row.get("rectification", "طلب تعديل درافت الفحص من شركة المعاينة الدولية"),
                }
                insp_discrepancies.append(item)
                all_discrepancies.append(item)
                if sev == "CRITICAL": total_critical += 1
                else: total_warning += 1

    insp_status = "WAIVED" if insp_is_waived else (
        "APPROVED" if (insp_available and inspection_review.status in ["Approved", "APPROVED", "Verified"]) else ("MODIFICATIONS_REQUESTED" if insp_discrepancies else ("REVIEW_PENDING" if insp_available else "NOT_STARTED"))
    )

    inspection_certificate = CentralArchiveDocumentSummary(
        document_type="INSPECTION_CERTIFICATE",
        title_ar="شهادة الفحص والتطابق (VoC / COC)",
        is_available=insp_available,
        is_mandatory=insp_is_mandatory,
        is_waived=insp_is_waived,
        waive_reason="الصنف غير خاضع لرقابة الصادرات والواردات (GOEIC) ولا يتطلب فحص ما قبل الشحن" if insp_is_waived else None,
        legal_requirement_note=insp_legal_note,
        status=insp_status,
        document_reference=inspection_review.certificate_number if inspection_review else None,
        details={
            "inspection_type": (inspection_review.inspection_type if inspection_review else "COC (Certificate of Conformity)"),
            "inspection_agency": (inspection_review.inspection_agency if inspection_review else "COTECNA / TÜV / SGS"),
            "certificate_number": inspection_review.certificate_number if inspection_review else "DRAFT-COC",
            "regulatory_authority": (inspection_review.regulatory_authority if inspection_review else "GOEIC"),
            "standard_specification": (inspection_review.standard_specification if inspection_review else None),
        },
        discrepancies=insp_discrepancies,
        raw_content=inspection_review.raw_input_text if inspection_review else None,
        last_updated=inspection_review.updated_at if inspection_review else None
    )

    # Dynamic Readiness Calculation (Only counts active non-waived mandatory docs)
    all_doc_summaries = [final_invoice, final_packing_list, draft_bl, certificate_of_origin, inspection_certificate]
    mandatory_docs = [d for d in all_doc_summaries if d.is_mandatory and not d.is_waived]
    req_count = len(mandatory_docs)
    approved_count = sum(1 for d in mandatory_docs if d.status == "APPROVED" or (d.is_available and len(d.discrepancies) == 0))

    base_score = (approved_count / req_count * 100.0) if req_count > 0 else 100.0
    deduction = (total_critical * 25.0) + (total_warning * 5.0)
    final_score = max(0.0, min(100.0, base_score - deduction))

    if total_critical == 0 and total_warning == 0 and approved_count == req_count:
        readiness_status = "READY_FOR_RELEASE"
        final_score = 100.0
    elif total_critical > 0 or total_warning > 0:
        readiness_status = "ACTION_REQUIRED"
    else:
        readiness_status = "IN_REVIEW"

    # Compile Import Requirements Summary
    import_requirements_summary = {
        "has_assessment": assessment is not None,
        "assessment_code": assessment.assessment_code if assessment else None,
        "hs_code": (assessment.hs_code if assessment else None) or "9403.10 / 9403.20",
        "commodity_description": (assessment.commodity_description if assessment else None) or "Office furniture & commercial goods",
        "country_of_origin": origin_country,
        "coo_required": coo_is_mandatory and not coo_is_waived,
        "coo_type": (assessment.coo_type if assessment else None) or ("EUR.1" if is_eu_origin else "Standard COO"),
        "coo_status": "Required" if (coo_is_mandatory and not coo_is_waived) else "Waived",
        "inspection_required": insp_is_mandatory and not insp_is_waived,
        "inspection_body": (assessment.inspection_body if assessment else None) or "COTECNA",
        "inspection_status": "Required" if (insp_is_mandatory and not insp_is_waived) else "Waived",
        "decree_43_applicable": assessment.decree_43_applicable if assessment else False,
        "white_list_verified": assessment.white_list_verified if assessment else False,
        "factory_registration_no": assessment.factory_registration_no if assessment else None,
        "import_permit_required": assessment.import_permit_required if assessment else False,
        "permit_issuing_authority": assessment.permit_issuing_authority if assessment else None,
        "overall_status": assessment.overall_status if assessment else "Completed / Inferred",
    }

    # Generate Supplier Email & WhatsApp text
    sys_acid = (acid_session.acid_number if acid_session else None) or imp_file.acid_number or "N/A"
    email_lines = [
        f"Subject: URGENT: Document Revision & Discrepancies Notice for Shipment {imp_file.import_file_code} (ACID: {sys_acid})\n",
        f"Dear {supplier.company_name if supplier else imp_file.supplier_name},\n",
        f"We have reviewed the draft shipping documents for Import File {imp_file.import_file_code} (Customs File: {imp_file.custom_file_number or 'N/A'}).",
        f"Please find below the required corrections and discrepancies that must be rectified before issuing the final original documents and uploading to CargoX:\n"
    ]
    if tariff_exemption_alert:
        email_lines.append(f"[IMPORTANT COMPLIANCE NOTE]: {tariff_exemption_alert}\n")
    if goeic_inspection_alert:
        email_lines.append(f"[INSPECTION NOTE]: {goeic_inspection_alert}\n")

    if all_discrepancies:
        for idx, item in enumerate(all_discrepancies, 1):
            sev_tag = "[CRITICAL / MANDATORY]" if item["severity"] == "CRITICAL" else "[WARNING / RECOMMENDATION]"
            email_lines.append(f"{idx}. {sev_tag} {item['document']}:")
            email_lines.append(f"   - Field / Item: {item['field']}")
            email_lines.append(f"   - Issue: {item['issue']}")
            email_lines.append(f"   - Required Correction: {item['rectification']}\n")
    else:
        email_lines.append("All reviewed documents are fully compliant with Egyptian Customs and Nafeza regulations. You may proceed to issue final originals.\n")

    email_lines.append(f"Consignee: {company.importer_name if company else imp_file.company_name}")
    email_lines.append(f"ACID Number: {sys_acid}")
    email_lines.append("\nBest regards,\nImport Operations Department - Sorour Logistics")
    email_text = "\n".join(email_lines)

    wa_lines = [
        f"*إشعار تعديلات المستندات - ملف الشحنة {imp_file.import_file_code}*",
        f"رقم القيد الجمركي ACID: `{sys_acid}`",
        f"المستورد: {company.importer_name if company else imp_file.company_name}",
        f"المورد: {supplier.company_name if supplier else imp_file.supplier_name}",
        "------------------------------------",
    ]
    if tariff_exemption_alert:
        wa_lines.append(f"🇪🇺 *تنبيه الإعفاء الجمركي:* يلزم تدوين `REVISED RULES` بالخانة 7 بشهادة EUR.1")
    if all_discrepancies:
        wa_lines.append("*التعديلات المطلوبة لإصدار الأصول:*")
        for idx, item in enumerate(all_discrepancies, 1):
            icon = "🔴" if item["severity"] == "CRITICAL" else "🟡"
            wa_lines.append(f"{icon} *{item['document']}* ({item['field']}):")
            wa_lines.append(f"   • المطلوب: {item['rectification']}")
    else:
        wa_lines.append("🟢 *كافة المستندات مطابقة بنسبة 100% وجاهزة لإصدار الأصول والرفع على كارجو إكس.*")
    wa_text = "\n".join(wa_lines)

    return CentralArchiveResponse(
        import_file_id=import_file_id,
        import_file_code=imp_file.import_file_code,
        custom_file_number=imp_file.custom_file_number,
        importer_name=(company.importer_name if company else imp_file.company_name) or "",
        supplier_name=(supplier.company_name if supplier else imp_file.supplier_name) or "",
        acid_number=sys_acid,
        port_of_loading=getattr(imp_file, 'port_of_loading', None) or getattr(imp_file, 'pol_name', None),
        port_of_discharge=getattr(imp_file, 'port_of_discharge', None) or getattr(imp_file, 'pod_name', None),
        total_packages=int(po_reconciliation.total_packages) if po_reconciliation else (booking.packages_count if booking else 0),
        total_gross_weight_kg=float(po_reconciliation.total_gross_weight_kg) if po_reconciliation else (booking.gross_weight if booking else 0.0),
        currency=po_reconciliation.currency if po_reconciliation else (getattr(imp_file, 'estimated_cost_currency', None) or "EUR"),
        fob_or_cif_amount=float(po_reconciliation.total_invoice_amount) if po_reconciliation else float(getattr(imp_file, 'estimated_cost', 0.0) or 0.0),
        readiness_status=readiness_status,
        readiness_score=round(final_score, 1),
        final_invoice=final_invoice,
        final_packing_list=final_packing_list,
        draft_bl=draft_bl,
        certificate_of_origin=certificate_of_origin,
        inspection_certificate=inspection_certificate,
        import_requirements_summary=import_requirements_summary,
        tariff_exemption_alert=tariff_exemption_alert,
        goeic_inspection_alert=goeic_inspection_alert,
        decree_43_alert=decree_43_alert,
        total_critical_discrepancies=total_critical,
        total_warning_discrepancies=total_warning,
        all_rectifications_checklist=all_discrepancies,
        supplier_email_rectification_text=email_text,
        supplier_whatsapp_rectification_text=wa_text,
    )
