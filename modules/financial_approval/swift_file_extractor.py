"""
Smart SWIFT & Bank Transfer Slip Multi-Format File Extractor
Supports: Word (.docx/.doc), Excel (.xlsx/.xls/.csv), PDF (.pdf), Images (.jpg/.jpeg/.png/.webp), and Text files.
"""

from __future__ import annotations

import io
import os
import re
import csv
import logging
from typing import Tuple, Optional

logger = logging.getLogger(__name__)

# Initialize RapidOCR engine once lazily
_rapid_ocr_engine = None


def get_ocr_engine():
    global _rapid_ocr_engine
    if _rapid_ocr_engine is None:
        try:
            from rapidocr_onnxruntime import RapidOCR
            _rapid_ocr_engine = RapidOCR()
        except Exception as e:
            logger.warning(f"Could not initialize RapidOCR: {e}")
            _rapid_ocr_engine = None
    return _rapid_ocr_engine


def normalize_swift_ocr_text(text: str) -> str:
    """
    Normalizes common OCR character misrecognitions in SWIFT MT103 and bank advice slips.
    """
    if not text:
        return ""

    # Replace fullwidth / Unicode colons and slashes
    text = text.replace('\uff1a', ':').replace('\uff0f', '/')

    # OCR 5OK -> 50K, 5OA -> 50A
    text = re.sub(r':?5[oO]([KA])', r':50\1', text)

    # Standardize tag lines e.g. "20/TRANSACTION..." or "59/Beneficiary..." to ":20:" or ":59:"
    standard_tags = ['20', '23B', '32A', '50K', '50A', '52A', '53A', '54A', '57A', '59', '59A', '70', '71A', '72']
    for tag in standard_tags:
        # Match tag at line start with optional leading colon, optional label, and optional table pipes
        text = re.sub(rf'(?m)^:?({tag})(?:/[^\n:|]+)?\s*[:\n|]\s*\|?\s*:?', rf':\1: ', text)

    return text


def extract_text_from_pdf(content_bytes: bytes) -> str:
    """Extracts text from PDF files using pdfplumber with pypdf and OCR fallbacks."""
    extracted_text = ""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
            pages_text = []
            for page in pdf.pages:
                t = page.extract_text(layout=True) or page.extract_text()
                if t and t.strip():
                    pages_text.append(t.strip())
            extracted_text = "\n\n".join(pages_text)
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}. Trying pypdf fallback.")
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content_bytes))
            parts = [p.extract_text() for p in reader.pages if p.extract_text()]
            extracted_text = "\n\n".join(parts)
        except Exception as e2:
            logger.warning(f"pypdf extraction failed: {e2}")

    # If PDF is scanned/image-only (less than 30 characters of text), try OCR on pages
    if len(extracted_text.strip()) < 30:
        try:
            import pypdfium2 as pdfium
            pdf = pdfium.PdfDocument(content_bytes)
            ocr_engine = get_ocr_engine()
            if ocr_engine:
                ocr_parts = []
                for i in range(len(pdf)):
                    page = pdf[i]
                    image = page.render(scale=2).to_pil()
                    img_byte_arr = io.BytesIO()
                    image.save(img_byte_arr, format='PNG')
                    ocr_res, _ = ocr_engine(img_byte_arr.getvalue())
                    if ocr_res:
                        ocr_parts.extend([line[1] for line in ocr_res])
                if ocr_parts:
                    extracted_text = "\n".join(ocr_parts)
        except Exception as ocr_err:
            logger.warning(f"PDF OCR extraction fallback failed: {ocr_err}")

    return extracted_text


def extract_text_from_word(content_bytes: bytes) -> str:
    """Extracts text and tables from Word (.docx / .doc) documents."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(content_bytes))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_vals = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_vals:
                    parts.append(" | ".join(row_vals))
        return "\n".join(parts)
    except Exception as e:
        logger.warning(f"Word docx extraction failed: {e}")
        return content_bytes.decode('utf-8', errors='ignore')


def extract_text_from_excel(content_bytes: bytes, filename: str) -> str:
    """Extracts text and table rows from Excel (.xlsx / .xls / .csv) files."""
    parts = []
    lower = filename.lower()
    if lower.endswith('.csv'):
        try:
            decoded = content_bytes.decode('utf-8-sig', errors='replace')
            reader = csv.reader(io.StringIO(decoded))
            for row in reader:
                cleaned = [c.strip() for c in row if c.strip()]
                if cleaned:
                    parts.append(" | ".join(cleaned))
            return "\n".join(parts)
        except Exception as e:
            logger.warning(f"CSV extraction failed: {e}")

    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content_bytes), data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"--- Sheet: {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if row_vals:
                    parts.append(" | ".join(row_vals))
        return "\n".join(parts)
    except Exception as e:
        logger.warning(f"Excel extraction failed: {e}")
        return content_bytes.decode('utf-8', errors='ignore')


def extract_text_from_image(content_bytes: bytes) -> str:
    """Extracts text from images (.jpg, .jpeg, .png, .webp, .bmp) using RapidOCR engine."""
    ocr_engine = get_ocr_engine()
    if not ocr_engine:
        return "OCR engine not available."
    try:
        result, _ = ocr_engine(content_bytes)
        if result:
            lines = [line[1] for line in result if line and len(line) > 1]
            return "\n".join(lines)
        return ""
    except Exception as e:
        logger.warning(f"Image OCR failed: {e}")
        return f"Image OCR Error: {e}"


def extract_text_from_swift_file(filename: str, content_bytes: bytes) -> Tuple[str, str]:
    """
    Main entry point for extracting text from any supported file format.
    Returns (extracted_raw_text, normalized_text).
    """
    lower = filename.lower()
    raw_text = ""

    if lower.endswith('.pdf'):
        raw_text = extract_text_from_pdf(content_bytes)
    elif lower.endswith(('.docx', '.doc')):
        raw_text = extract_text_from_word(content_bytes)
    elif lower.endswith(('.xlsx', '.xls', '.csv')):
        raw_text = extract_text_from_excel(content_bytes, filename)
    elif lower.endswith(('.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff', '.tif')):
        raw_text = extract_text_from_image(content_bytes)
    else:
        # Plain text / fallback
        try:
            raw_text = content_bytes.decode('utf-8')
        except UnicodeDecodeError:
            raw_text = content_bytes.decode('latin-1', errors='ignore')

    normalized = normalize_swift_ocr_text(raw_text)
    return raw_text, normalized
